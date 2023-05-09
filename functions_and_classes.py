#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# pylint: disable=line-too-long
"""
Created on Wed Nov 25 23:46:05 2020.

@author: ena
"""

# %% Imports

# Basics
import rpy2.robjects.packages as rpackages
from rpy2 import robjects
from matplotlib.backends.backend_pdf import PdfPages
# import matplotlib.cm as cm
import matplotlib.pyplot as plt
import matplotlib
import seaborn as sb
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
# from docx import Document as Document_compose
# from docxcompose.composer import Composer
from docx.shared import Pt
from docx.shared import Inches
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.section import WD_ORIENT
import docx
import functools
import numpy as np
import pandas as pd
import scipy as sp
import re
import os
import warnings
import copy
from functions_analysis import corr_dfs, try_regress, descriptives_tables, compare_data
# from functions_visualization import cap_title, square_grid, heatmaps, square_heatmap, plot_kde, cap
from functions_visualization import cap_title, square_grid, heatmaps, square_heatmap
from functions_documentation import string_list, combine_word_documents, results_section_documents, write_table
from functions_data_cleaning import type_convert, check_frequencies, wh, dict_part, try_float_format
from functions_mplus import read_savedata, read_fit, read_loadings, cfa_syntax

# R-Python Interface
R_LIBS = robjects.r('.libPaths()')[0]
base = rpackages.importr('base')
utils = rpackages.importr('utils')
stats = rpackages.importr('stats')
mplus = rpackages.importr('MplusAutomation', lib_loc=R_LIBS)
# from rpy2.robjects import pandas2ri
# pandas2ri.activate()

# Options
pd.options.display.max_columns = 100
pd.options.display.max_rows = 200
pd.options.display.large_repr = 'truncate'


# %% Basic Custom Functions

def cap(string):
    """Capitalize the first letter of a string."""
    return string[0].upper() + string[1:]


def colorbar_min_range(data, cbar_min_range=[-1, 1]):
    """Return specified range if encompasses data; otherwise, range of data."""
    mini = data.min().min() if type(data) == pd.DataFrame else min(data)
    maxi = data.max().max() if type(data) == pd.DataFrame else max(data)
    cmin = cbar_min_range[0] if mini > cbar_min_range[0] else mini  # color bar minimum
    cmax = cbar_min_range[1] if maxi < cbar_min_range[1] else maxi  # color bar maximum
    return cmin, cmax


def model_dictionary_to_content(model_names, model_dictionary, fac_series=None):
    """Make series with list for each model of lists of factor indicators."""
    all_indicators = [key for key in model_dictionary.keys()]  # all factor indicators appearing in models
    model_content = pd.DataFrame(model_dictionary, index=[model_names]).T.replace(np.nan, '')
    model_content.columns = model_names
    if fac_series is None:
        fac_series = model_content.apply(lambda f: pd.unique(f)).apply(lambda x: x[x != ''])
    fac_content = model_content.apply(lambda df: [df.index.values[df == f] for f in fac_series[df.name]])
    indicator_factors = model_content.apply(lambda m: pd.unique(np.array(m[m != '']).flatten()), axis=1)
    return model_content, fac_content, fac_series, indicator_factors, all_indicators


def make_model_info(model_dictionary=None, model_content=None, fac_series=None, indicator_dictionary=None,
                    model_names=None, categorical=None):
    """Make information about models in various formats (e.g., series, dictionaries, dataframes)."""
    if (model_content is None) and (model_dictionary is None):
        arg_1 = 'dataframe of factor labels with model names as columns & indicators as row indices (model_content)'
        arg_2 = 'dictionary with lists of factor lables for each model & indicators as keys (model_dictionary)'
        raise Exception('Either a %s or %s must be provided.' % (arg_1, arg_2))
    if model_names is None:
        if model_content is not None:
            model_names = list(model_content.columns)
        elif fac_series is not None:
            model_names = list(fac_series.index)
        else:
            raise Exception('Either model_content or fac_series must be specified if model names are not provided.')
    if model_dictionary is None:
        df = model_content.T.reset_index(drop=True)  # transpose DF & drop model names as indices
        model_dictionary = dict(zip(df.columns, [list(df[c]) for c in df.columns]))  # model dictionary (keys=indicators)
    if model_content is None:
        if model_dictionary is not None:
            model_content = pd.DataFrame(model_dictionary, index=model_names).T
    if fac_series is None:
        fac_series = model_content.apply(lambda f: pd.unique(f)).apply(lambda x: x[x != ''])
    model_content = model_content.replace(np.nan, '')
    fc = model_content.apply(lambda df: [df.index.values[df == f] for f in fac_series[df.name]])
    fac_content = model_content.apply(lambda x: dict(zip(dict(zip(model_names, fac_series))[x.name], fc[x.name])),
                                      result_type='reduce')
    if indicator_dictionary is not None:
        models = pd.concat([pd.DataFrame(pd.Series(indicator_dictionary), columns=['']),
                            pd.DataFrame(model_dictionary, index=model_names).T], axis=1)
    else:
        models = model_content
    num_factors = [len(fac_series[m]) for m in model_names]
    multi_ix = pd.MultiIndex.from_tuples([tuple([str(i[0]), i[1]]) for i in zip(num_factors, model_names)])  # multi-ix
    desc = pd.DataFrame(model_content, columns=multi_ix).replace(np.nan, '')  # content: # factors & models as multi-index
    num_factors = [len(fac_series[m][0]) for m in model_names]
    all_indicators = [key for key in model_dictionary.keys()]  # all factor indicators appearing in models
    mod_content = pd.DataFrame(model_dictionary, index=model_names).T.astype(str)
    indicator_factors = model_content.apply(lambda m: pd.unique(np.array(m[m != '']).flatten()), axis=1)
    out = {'model_names': model_names, 'model_dictionary': model_dictionary, 'descriptives_models': desc,
           'fac_content': fac_content, 'all_indicators': all_indicators, 'num_factors': num_factors,
           'indicator_factors': indicator_factors, 'categorical': categorical, 'fac_series': fac_series,
           'model_content': mod_content, 'models': models}
    return out


def make_file_info(file, models, analysis, suff, cap_suff='', directory=None, dir_file=None):
    """Make a dictionary of file names based on model characteristics."""
    if directory is None:
        directory = os.getcwd()
    if dir_file is None:
        dir_file = str(re.sub('.csv', '', file) + '_mplus.csv')
    # if directory not in dir_file: dir_file = '%s/%s'%(directory, dir_file)
    if suff not in [0, None]:  # if has extra suffix for .inp, .out, and .csv file stems
        stem = [m + '_%s_%s' % (analysis, suff) for m in models]  # stem
    else:
        stem = [m + '_' + analysis for m in models]  # file stem
    out = [os.getcwd() + '/%s/' % directory + m for m in stem]  # MPlus syntax & output file paths
    if cap_suff != '':
        cap_suff = ' %s' % cap_suff if cap_suff[0] != ' ' else cap_suff  # add caption suffix space if need
    files = {'mplus_directory': directory, 'file_stem': stem, 'data_file': file, 'file_suffix': suff,
             'analysis': analysis, 'caption_suffix': cap_suff, 'mplus_out_files': out, 'data_mplus_file': dir_file}
    return files


def color_tab(x, sig=None):
    """Color based on sign & (optionally) significance."""
    y = 0 if (sig is not None) and (str(x).count('*') < sig.count('*')) else try_float_format(x, ignore='[*]')
    if type(y) is not float:
        return 'color: black'
    else:
        return 'color: %s' % ('black' if y == 0 else ['green', 'red'][y < 0])


# %% Base & Factor Model Classes

class Factor_Model:
    """
    Model Contents Class.

    Arguments
    ---------
    data_file : str
        Path to source data
    mplus_directory : str
        Directory in which to write MPlus data, syntax, & output
    data_file_mplus : str
        File name to which to write MPlus-friendly version of data
    model_dictionary : list of dictionaries
        (Keys = factors, values = indicators) with content for each model (see example below)
    model_content : pandas.DataFrame
        Dataframe version of above (models = columns, rows = indicators, cells = factors on which indicators load)
    indicator_dictionary : dict, optional
        Dictionary, keyed by indicators, with strings describing the indicators
    cap_indicator_table : str, optional
        Title for indicator description table
    analysis : str, optional (default CFA)
        CFA or Bifactor
    model_names : list
        Model names
    data_file : str
        Name of file containing MPlus-friendly data (full file path)
    data_columns: list
        List of column names for MPlus-friendly data
    categorical : bool or list
        True if all categorical indicators, False if none, or list of strings identifying categorical indicators
    file_suffix : str
        Additional suffix for model (if more than model name and analysis desired)
    index_col : str
        Name of index column(s)
    subpopulation : str, optional
        String with line of MPlus code to indicate subpopulation for factor analysis (e.g. AGE > 30)
    kwargs_read_csv: optional
        Extra keyword arguments to pass to pandas.read_csv() when loading data file

    Methods
    -------
    create_syntax() : Create attribute with MPlus CFA syntax
    write_syntax() : Write syntax to file
    run_mplus() : Run syntax created by write_syntax() method
    read_mplus() : Extract factor scores & loadings
    factor_correlations() : Extract factor correlations within-models & across models
        (comparison to common base model)
    analyze_correlations() : Compute correlations between factors & outcome variables
    analyze_regressions() : Perform regressions with factors as IVs
    write_results() : Construct HTML & LaTeX syntax for results tables
    heatmap() : Display and/or save heat maps for correlations (among factors, between factors & DVs) & regressions
        Factor correlation heat maps display magnitudes of correlation coefficients
        Correlation heat maps display magnitudes of correlation coefficients (with or without significance stars)
        Regression heat maps display -1, 0, 1 values representing -, non-significant, & + effects, respectively

    Major Attributes
    ----------------
    fac_series : Series (indexed by model names) with lists of models' factors
    fac_content : Series (indexed by model names) with dictionaries of lists of indicators for each model factor
    results_<fit, correlation, regression, r2> : Fit statistics, correlations, & regression & r-squared results
        (available after running read_mplus(), analyze_correlations(), & analyze_regressions(), respectively)
    results_<...>.<latex, html> : HTML & LaTeX syntax for results tables
        Syntax is stored as attributes of results tables & is thus accessed by appending a .html or .latex.
        N.B. For lists, you must access list elements to access .html & .latex attributes
    output() : MPlus output (as R object; available after running read_mplus())

    """

    def __init__(self, data_file, index_col='', data_file_mplus=None, data_columns=None, subpopulation=None,
                 data_types=None, kwargs_read_csv={'keep_default_na': False}, recode_map={}, na_flag='*',
                 indicator_dictionary=None, check_dict=None, reverse_code_flag='_r', model_dictionary=None,
                 cap_indicator_table='AUDADIS-5 Items Corresponding to DSM–5 Symptoms', indicator_items=None,
                 model_content=None, model_names=None, file_suffix=0, fac_series=None, analysis='CFA',
                 categorical=True, mplus_directory='MPlus', caption_suffix='', auxiliary=True, **kwargs):
        """Initialize object from Factor_Model Class."""
        if mplus_directory not in os.listdir():
            os.system('mkdir %s' % mplus_directory)  # make MPlus directory if absent
        model_info = make_model_info(model_dictionary=model_dictionary, categorical=categorical,
                                     fac_series=fac_series, model_content=model_content, model_names=model_names,
                                     indicator_dictionary=indicator_dictionary)
        if model_names is None:
            model_names = model_info['model_names']
        file_info = make_file_info(data_file, model_names, analysis, file_suffix, cap_suff=caption_suffix,
                                   directory=mplus_directory, dir_file=data_file)
        info = {**model_info, **file_info}
        self.indicator_dictionary = indicator_dictionary
        self.indicators = pd.DataFrame(pd.Series(indicator_items), columns=['Item']).rename_axis('Indicator')
        [setattr(self, f, info[f]) for f in info]  # make attributes
        self.gather_attributes('file', 'files')  # gather file-related attributes
        self.data_original = pd.read_csv(data_file, low_memory=False,
                                         # names=data_columns,
                                         **kwargs_read_csv)
        if data_types is not None:
            self.data_original = type_convert(self.data_original, data_types)  # type-conversion
        if subpopulation is not None:  # if subpopulation is specified...
            self.data_original = self.data_original[eval('self.data_original.%s' % subpopulation)]  # sub-sample
        self.captions = {'models': {}}  # initialize a dictionary to hold captions later
        self.captions.update({'models': 'Factors and Content of Competing Models',
                              'indicators': cap_indicator_table})
        self.data_info = {'file': data_file, 'file_mplus': data_file_mplus, 'mplus_directory': mplus_directory,
                          'working_directory': os.getcwd(), 'columns': list(self.data_original.columns),
                          'index_col': index_col, 'na_flag': na_flag, 'reverse_code_flag': reverse_code_flag,
                          'data_types': data_types, 'auxiliary': auxiliary, 'subpopulation': subpopulation,
                          'recode_map': recode_map}
        self.check_dict = check_dict
        self.data_check = {'frequencies': []}
        if self.check_dict is not None:  # check frequencies (if desired)
            self.data_check['frequencies'], _ = check_frequencies(self.data_original, self.check_dict, plot=False,
                                                                  rev_code_flag=self.data_info['reverse_code_flag'])
        self.syntax = None

    def gather_attributes(self, keyword, new_name=None):
        """Gather attributes into dictionary by keyword & delete individual."""
        if new_name is None:
            new_name = keyword  # use keyword as new attribute name if not specified
        old_names = wh([keyword in d for d in dir(self)], True, dir(self))  # names of file name-type attributes
        attrs = [self.__getattribute__(a) for a in old_names]  # values of those attributes
        [self.__delattr__(a) for a in old_names]
        self.__setattr__(new_name, dict(zip(old_names, attrs)))  # gather keyword-matching attributes into 1 dictionary
        if 'attributes_old' not in dir(self):
            self.attributes_old = {}  # initialize old attributes attribute if needed
        self.attributes_old.update({new_name: old_names})  # store old attribute names keyed by dictionary attribute name

    def restore_old_attributes(self):
        """Restore individual attributes from gathered dictionaries."""
        if 'attributes_old' in dir(self):
            [self.__setattr__(a, self.files[a]) for a in self.files]  # restore individual file name-type attributes
        else:
            print('No old attributes to restore.')

    def get_attributes_from_stems(self, stems, **kwargs):
        """Return all attributes that have a given string in their name."""
        if ((type(stems) is list) or (type(stems) is np.ndarray)) and (len(stems) > 1):  # if multiple stems
            out = [np.nan] * len(stems)
            for s in range(len(stems)):  # iterate through stems
                out[s] = [self.__getattribute__(a) for a in wh([stems[s] in x for x in dir(self)], 1, dir(self))]
        else:  # if only one stem provided
            out = [self.__getattribute__(a) for a in wh([s in x for x in dir(self)], 1, dir(self))]
        print(out)
        return out

    def clear_attributes(self, attributes=['descriptives', 'results', 'figures'], **kwargs):
        """Clear Attributes related to certain categories."""
        for att in attributes:  # for all attribute types specified
            for a in wh(['%s_' % att in x for x in dir(self)], True, dir(self)):  # for all attributes with that stem
                if a in dir(self):
                    self.__delattr__(a)

    def create_syntax(self, write=True, standardization='standardized', free_load=True, op='Linux',
                      cores=None, threads=None, starts=None, stit=None, ESTIMATOR='wlsmv', faster=False,
                      cluster=None, stratification=None, weight=None, TYPE=None, montecarlo=None,
                      variable_statement_extras=None, analysis_statement_extras=None, auxiliary=True,
                      model_statement_extras=None, output_statement_extras=None, directory_agnostic=False,
                      plot_statement_extras=None, save_statement_extras=None, run=False,
                      interactive=None, **kwargs):
        r"""
        Write MPlus syntax.

        Arguments
        ---------
        write : bool
            Write syntax? (If false, will just store as attribute)
        standardization : str
            MPlus standardization option (e.g. 'standardized', 'stdyx', 'stdx', 'stdy')
        free_load = bool,
            Free 1st indicators' loadings & fix factor variances to 1? False: Free factor variances, fix 1st loadings
        stratification : str, optional
            Name of sample stratification variable
        cluster : str, optional
            Name of sample clustering variable
        weights : str, optional
            Name of sample weights variable
        montecarlo : int, optional
            Number of montecarlo integrations
            (500 or 1000 runs faster than default...check to ensure no negative ABSCHANGE values in TECH8)
        op : str, optional
            Operating system. 'Linux' to create .sh exectuables to run MPlus batches. 'Windows' for .bat.
        TYPE : str, optional
            Type command under the analysis statement (e.g. 'complex')
        <variable, analysis, model, output, plot, save>_statement_extras : str, optional
            String (or list of strings for each model for analysis & model)
                with extra syntax to append to the relevant statements. First new line & tab already included,
                if more than one line, include \n\t between them, and be sure to end all statements with a semicolon.
        auxiliary : bool, optional
            Whether to include other data columns in MPlus-returned "savedata." The default is True.
        **kwargs : str, optional
            Other keyword arguments corresponding to MPlus options, formatted as in MPlus (e.g. ESTIMATOR='MLR').
        """
        self.data_info.update({'mplus_directory_agnostic': directory_agnostic, 'auxiliary': auxiliary})
        self.syntax = [np.nan for m in range(len(self.model_names))]
        self.data_original.to_csv(self.data_info['file_mplus'], index=False, header=False)  # write MPlus-friendly file
        self.mplus_data = pd.read_csv(self.data_info['file_mplus'],
                                      index_col=False, names=self.data_original.columns)  # read csv as it was written
        self.mplus_arguments = {'standardization': standardization, 'free_load': free_load, 'ESTIMATOR': ESTIMATOR,
                                'cores': cores, 'threads': threads, 'starts': starts, 'stit': stit,
                                'cluster': cluster, 'stratification': stratification, 'weight': weight,
                                'subpopulation': self.data_info['subpopulation'], 'TYPE': TYPE,
                                'auxiliary': auxiliary, 'montecarlo': montecarlo,
                                # 'analysis_statement_extras': analysis_statement_extras,
                                # 'model_statement_extras': model_statement_extras,
                                'output_statement_extras': output_statement_extras,
                                'plot_statement_extras': plot_statement_extras,
                                'variable_statement_extras': variable_statement_extras,
                                'save_statement_extras': save_statement_extras,
                                'data_column_names': self.data_original.columns}
        if 'mplus_out_files' not in dir(self):
            self.restore_old_attributes()
        out_files = self.files['mplus_out_files']  # output file names
        if type(analysis_statement_extras) is not list:
            analysis_statement_extras = [analysis_statement_extras] * len(self.model_names)
        if type(model_statement_extras) is not list:
            model_statement_extras = [model_statement_extras] * len(self.model_names)
        for m in range(len(self.model_names)):
            a_extras, m_extras = analysis_statement_extras[m], model_statement_extras[m]  # extras
            dir_ag = re.sub('.*/%s/' % self.data_info['mplus_directory'], '', out_files[m])  # directory-less
            dir_ag_df = re.sub('.*/%s/' % self.data_info['mplus_directory'], '', self.data_info['file_mplus'])  # ""
            out_file = dir_ag if directory_agnostic else out_files[m]  # MPlus output files (w/o .out)
            data_file = dir_ag_df if directory_agnostic else self.data_info['file_mplus']  # MPlus data file
            self.syntax[m] = cfa_syntax(self.model_names[m], self.fac_series[m],
                                        self.fac_content[self.model_names[m]], out_file=out_file, preview=True,
                                        na_flag=self.data_info['na_flag'], analysis=self.analysis, plots=True,
                                        write=False, faster=faster, data_file=data_file,
                                        categorical=self.categorical, id_variable=self.data_info['index_col'],
                                        analysis_statement_extras=a_extras, model_statement_extras=m_extras,
                                        interactive=interactive, **self.mplus_arguments)
            if write:
                print('\n\nWriting %s syntax to %s' % (self.model_names[m], str(out_files[m] + '.inp')))
                script_file = open(str(out_files[m] + '.inp'), 'w')
                script_file.writelines(self.syntax[m])
                script_file.close()
        print('\n\n=================\n' + self.model_names[m] + '\n=================\n\n' + self.syntax[m])  # example
        if write:
            batch_i = (self.data_info['mplus_directory'], re.sub(self.model_names[0], '', self.files['file_stem'][0]))
            exe_file = str('%s/batch%s.sh' % batch_i) if op.upper() == 'LINUX' else str('%s/batch%s.bat' % batch_i)
            print('\n\nWriting Batch File to %s' % exe_file)
            b = open(exe_file, 'w')
            if op.upper() == 'LINUX':  # for Linux batch run executable
                outs = [re.sub('.*/%s/' % self.data_info['mplus_directory'], '', i) for i in out_files]
                header, cmnds = '#!/bin/bash\n\n', ['mplus %s.inp > %s_log.txt 2>&1 &' % (i, i) for i in outs]
                body = functools.reduce(lambda x, y: '%s\n%s' % (x, y), cmnds)
                b.writelines(header + body)
                os.system('chmod +x "%s"' % exe_file)  # make executable
            else:  # for Windows batch
                b.writelines(functools.reduce(lambda x, y: '%s.inp\n%s.inp' % (x, y), out_files))
            b.close()
        if run:
            for m in out_files:
                os.system('mplus %s.inp %s.out' % (m, m))

    def read_syntax(self, out_files=None, **kwargs):
        """Read MPlus input."""
        self.syntax = {}
        if out_files is None:
            out_files = self.files['mplus_out_files']  # output file names
        for i, m in enumerate(self.model_names):
            print('\n\nReading %s syntax from %s' % (m, str(out_files[i] + '.inp')))
            script_file = open(str(out_files[i] + '.inp'), 'rb')
            syntax = script_file.readlines()
            script_file.close()
            self.syntax.update({m: syntax})

    def read_mplus(self, check=True, join_data=False, **kwargs):
        """Read MPlus output."""
        # Factor Scores
        if 'mplus_out_files' not in dir(self):
            self.restore_old_attributes()
        ix, na_flag = self.data_info['index_col'], self.data_info['na_flag']
        args = [[i[0], ix, i[1]] for i in zip(self.fac_series, self.files['mplus_out_files'])]  # read_savedata
        out = [read_savedata(self.data_original, *a, join_data=join_data) for a in args]  # extract MPlus output
        df, self.mplus_output = [o[0] for o in out], [o[1] for o in out]  # unpack savedata/output
        self.data = dict(zip(self.model_names, df))  # concatenate such that models are super-columns
        compare = []  # initialize list for MPlus savedata-original data check/comparisons
        if check:
            do = self.data_original.set_index(ix) if ix in self.data_original.columns else self.data_original
            for m in self.model_names:  # iterate across models to check alignment of MPlus savedata & original
                try:
                    sd, factors = self.data[m], self.fac_series[m]  # savedata & factors
                    if sd[factors].replace(na_flag, np.nan).dropna().empty:
                        co = 'Failed factor extraction'
                        continue  # skip failed extractions
                    ix = sd.index.intersection(do.index)  # index intersection with original data
                    cx = sd.columns.intersection(do.columns)  # column intersection with original data
                    co = sd.loc[ix][cx].compare(do.loc[ix][cx].replace(na_flag, np.nan))
                    if co.empty is False:
                        for y in list(pd.unique([c[0] for c in co.dropna(how='all').columns])):
                            try:  # try converting to float to see if only mismatch because read as string
                                floated = co[y].replace(na_flag, np.nan).astype(float)
                                if all(floated.apply(lambda x: all(pd.isnull(x)) or x[0] == x[1], axis=1)):
                                    self.data[m].loc[:, y] = self.data[m][y].astype(float)
                                    do.loc[:, y] = do.replace(na_flag, np.nan).astype(float)
                            except Exception as err:
                                print(err, '\n\nFailed to convert to float.\n\n')
                        co = self.data[m].loc[ix][cx].compare(do.loc[ix][cx].replace(na_flag, np.nan))  # re-check
                        # mr = ['%s Correlation: %d'%(y, sp.stats.pearsonr(co[y], do[y])[0]) for y in mis] # correlation
                        info = ('*' * 80, m, str(list(pd.unique([c[0] for c in co.dropna(how='all').columns]))), '*' * 80)
                        warnings.warn('\n\n%s\n\n%s Original-MPlus data mismatch: %s\n\n%s' % info)
                    compare = compare + [co]
                except Exception as err:
                    warnings.warn(f'{err}/n/nCould not check savedata for model {m}.')
        if 'descriptives_factors' not in self.captions.keys():
            self.captions.update({'descriptives_factors': {}})
        cap_stems = ['%s Summary of Model Fit Statistics', 'Standardized %s Factor Loadings']
        caps = ['%s%s' % (i % self.analysis, self.caption_suffix) for i in cap_stems]
        self.captions['descriptives_factors'].update({'loadings': caps[1]})  # update captions
        self.captions.update({'factor_loadings': caps[1]})  # update captions
        self.captions.update({'results_fit': caps[0]})  # update captions
        warns = dict(zip(self.model_names, [wh(o.names, 'warnings', o) for o in self.mplus_output]))
        self.mplus_warnings = ['\n\n%s\n%s Warnings:\n%s\n\n%s' % ('*' * 18, o, '*' * 18,
                                                                   str(warns[o])) for o in warns]
        self.data_check.update({'mplus': compare})
        # replace = [dict(zip(d.columns, [{self.data_info['na_flag']: np.nan}]*d.shape[1])) for d in self.data]
        self.data = dict([(d, self.data[d].replace(na_flag, np.nan)) for d in self.data])
        # Fit Statistics & Loadings
        self.results_fit = read_fit(self.model_names, mplus_output=self.mplus_output, digits=3)
        try:
            self.results_fit = pd.DataFrame(self.results_fit)
            if 'Parameters' in self.results_fit.columns:
                if '# of Parameters' in self.results_fit.columns:
                    self.results_fit = self.results_fit.drop('Parameters', axis=1)
        except Exception as err:
            print(err, '\n\nCould not create results_fit dataframe.')
        lds = read_loadings(self.fac_series, analysis=self.analysis, mplus_output=self.mplus_output,
                            all_indicators=self.all_indicators)
        self.descriptives_factors = {'loadings': lds}
        self.factor_loadings = lds
        self.data_check.update({'mplus': compare})  # unpack output

    def factor_correlations(self, i=0, digits=3,
                            plot=True, plot_location=None, orient='landscape', **kwargs):
        """
        Store & plot factor Correlations (within-model + between-models).

        Arguments
        ---------
        i : int, optional
            Numerical index of "base" model for the inter-model comparison (position in self.model_names)
        digits : int, optional
            Number of digits to after decimal point to which to round/to include
        """
        self.base_model = self.model_names[i]  # store base model for comparison
        facs = [self.data[self.model_names[f]][self.fac_series[f]] for f in range(len(self.fac_series))]
        mats = [[corr_dfs(fs, df_2=d, digits=digits) for fs in facs] for d in [None, facs[i]]]  # factor correlations
        within, wi_p = [pd.concat([r[x] for r in mats[0]], keys=self.model_names) for x in [0, 1]]  # within
        between, bt_p = [pd.concat([r[x] for r in mats[1]], keys=self.model_names) for x in [0, 1]]  # between
        t, tp = [pd.concat(x, axis=1, keys=['Within', 'Between']) for x in [[within, between], [wi_p, bt_p]]]  # both
        wis = pd.concat([r[0] for r in mats[0]], keys=self.model_names)
        tab_wi = wis.applymap(lambda x: '' if pd.isnull(x) else ('{:.%df}' % digits).format(x))  # formatted within
        tab_bt = between.applymap(lambda x: '' if pd.isnull(x) else ('{:.%df}' % digits).format(x))  # formatted between
        dictny = {'correlations': t, 'correlations_p': tp, 'correlations_within': [r[0] for r in mats[0]],
                  'correlations_within_p': [r[0] for r in mats[0]], 'correlations_between': between,
                  'correlations_between_p': bt_p, 'correlations_within_table_2': tab_wi,
                  'correlations_between_table': tab_bt}
        self.descriptives_factors.update(dictny)
        bt_suff = [' with %s %s Model Factors' % (self.model_names[i], self.analysis)] * 4
        caps = ['Correlations%s' % i for i in [''] * 2 + [f' of factors within {self.analysis} Models'] * 2 + bt_suff]
        caps = [c + self.caption_suffix for c in caps]  # add object's caption suffix
        if 'captions' not in dir(self):
            self.captions = dict()
        if 'descriptives_factors' not in list(self.captions.keys()):
            self.captions.update({'descriptives_factors': {}})  # initialize factor descriptives dictionary in captions
        self.captions['descriptives_factors'].update(**dict(zip(list(dictny.keys()), caps)))  # update captions
        wia = [x.T.rename_axis('').reset_index().T for x in self.descriptives_factors['correlations_within']]
        wia = pd.concat(wia, keys=self.model_names).replace(np.nan, '')
        self.Factor_Correlations = {'Within-Model': wia, 'Between-Model': tab_bt}
        base_info = (self.analysis, self.model_names[i], self.caption_suffix)
        win = (self.analysis, self.caption_suffix)
        self.captions.update({'Factor_Correlations':
                              {'Within-Model': '%s Factor Inter-Correlations Within-Model%s' % win,
                               'Between-Model': '%s Factor Correlations with %s Model Factors%s' % base_info}})
        factors = self.fac_series
        if self.analysis.lower() == 'bifactor':
            factors = [['GF'] + m for m in factors]  # add general factor (if needed)
        all_facs = [self.data[self.model_names[m]].loc[:, factors[m]] for m in range(len(self.model_names))]
        self.factor_scores = pd.concat(all_facs, axis=1, keys=self.model_names)  # all model factor scores in one
        all_corrs = self.factor_scores.corr()  # all correlations (across models)
        if all(all_corrs >= 0):  # color bar normalization if all positive
            vmin, vmax, cmap = min(min(all_corrs.min()), 0), max(max(all_corrs.max()), 1), 'Reds'
        else:  # color bar normalization if NOT all positive
            vmin, vmax, cmap = min(min(all_corrs.min()), -0.5), max(max(all_corrs.max()), 1), 'coolwarm'
        self.correlations_factors_all = all_corrs.applymap(lambda x: ('{:.%df}' % digits).format(x))  # format
        for r in range(all_corrs.shape[0]):
            for c in range(all_corrs.shape[1]):
                if (r == c) or (r < c):
                    all_corrs.iloc[r, c] = np.nan  # mask diagnoal & upper triangle
                    self.correlations_factors_all = ''  # blank diagnoal & upper triangle
        all_corrs.columns = pd.MultiIndex.from_tuples([(re.sub('_', ' ', c[0]), c[1]) for c in all_corrs.columns])
        matplotlib.rc('font', **{'family': 'serif', 'size': 10})
        fig, axes = plt.subplots(figsize=(35, 20))  # initialize figure
        sb.heatmap(all_corrs, cmap=cmap, vmin=vmin, vmax=vmax, ax=axes,
                   xticklabels=[(re.sub('_', ' ', c[0]), c[1]) for c in all_corrs.columns],
                   yticklabels=[(re.sub('_', ' ', c[0]), c[1]) for c in all_corrs.columns])  # heatmap
        axes.set_xlabel('')
        axes.set_ylabel('')
        fig.tight_layout()
        if 'Plots' not in os.listdir():
            os.system('mkdir Plots')  # make plots folder if absent
        fig_file = 'Plots/correlations_factors_all_%s' % self.analysis
        suf = self.files['file_suffix'] if 'file_suffix' in self.files.keys() else self.file_suffix
        if suf != 0:
            fig_file = fig_file + '_' + suf
        fig.savefig(fig_file + '.jpeg')
        return fig_file, orient


# %% Analysis Class

class Factor_Model_Analysis(Factor_Model):
    """Analysis."""

    # Initialization Method
    def __init__(self, DV_subsets_dictionary=None, binary=None, base_model=0, join_data=False,
                 cap_descriptives='Descriptive Statistics for',
                 syntax='create', run_mplus_syntax=False, **kwargs):
        """Initialize analysis."""
        data_file = kwargs.pop('data_file')
        super(Factor_Model_Analysis, self).__init__(data_file, **kwargs)  # initialize parent
        # Syntax & Output
        if syntax == 'create':
            self.create_syntax(run=run_mplus_syntax, **kwargs)  # create MPlus syntax
        else:
            self.read_syntax(run=run_mplus_syntax, **kwargs)
        self.read_mplus(join_data=join_data)  # read output
        for m in self.model_names:
            columns = list(self.data[m].columns)
            for i, c in enumerate(columns):  # drop duplicate columns
                if self.data[m][[c]].shape[1] > 1:
                    warnings.warn('\n\n%s\n\nDuplicate columns in data (%s). Re-naming.\n\n%s'
                                  % ('=' * 80, c, '=' * 80))
                    columns[i] = c + '_%d' % i if wh(columns, c)[0] < i else c  # re-name if not 1st of duplicates
        self.base_model = base_model  # base model for between-model factor correlations
        if 'files' not in dir(self):
            self.files = {'figures': {}, 'results': {}}
        if 'figures' not in self.files.keys():
            self.files.update({'figures': {}})
        if 'captions' not in dir(self):
            self.captions = dict()
        try:
            fig_file, orient = self.factor_correlations(i=self.base_model)
        except Exception as err:
            print(err)
            print('\n\n\nCould not obtain factor correlations.')
        try:
            self.files['figures'].update({'main': {'correlations_factors_all': (fig_file + '.jpeg', orient)}})
            self.captions.update({'correlations_factors_all': 'Model Factor Correlations'})  # update captions
        except Exception as err:
            print(err)
            print('\n\n\nCould not store file names and/or captions.')
        # Remove Invariable Variables from DVs
        self.DV_subsets_dictionary = DV_subsets_dictionary
        DV_dict = copy.deepcopy(DV_subsets_dictionary)  # DVs
        for k in DV_dict:
            try:
                invars = np.array(wh([self.data[self.model_names[0]][y].var() == 0 for y in DV_dict[k]], 1, DV_dict[k]))
                if len(invars.ravel()) > 0:
                    for x in invars.ravel():  # iterate invariables
                        self.DV_subsets_dictionary[k].remove(x)  # remove from DV dictionary
            except Exception as err:
                warnings.warn(f'{err}\n\nCould not check {k} outcomes to see if invariable.')
        if binary is None:
            binary = [False] * len(self.DV_subsets_dictionary.keys())  # assume not binary if not specified
        self.binary = binary
        # Check All Models Same DV Data
        try:
            compare = []  # empty list for checks
            DVs_all = pd.Series(copy.deepcopy(DV_subsets_dictionary)).explode()
            for m in self.model_names[1:]:  # compare DV values for 1st model's dataframe to subsequent models
                comp = self.data[self.model_names[0]][DVs_all].compare(self.data[m][DVs_all])
                if comp.empty is False:
                    warnings.warn('DV mismatch across models\' datasets!')
                compare = compare + [comp]
            self.data_check.update({'model_DVs': compare})
        except Exception as e:
            print(e)
            warnings.warn('\n\n%s\n\nFailed to conduct check for equivalency of models\' outcome data!\n\n%s'
                          % ('*' * 80, '*' * 80))
        self.descriptives_groups = None

    def analyze_descriptives(self, descriptives_groups=None, descriptives_groups_rename=None,
                             DVs_all=None, **kwargs):
        """Obtain outcome descriptives."""
        self.descriptives_groups = descriptives_groups
        if DVs_all is None:
            DVs_all = pd.Series(self.DV_subsets_dictionary).explode()
        try:
            desc_vars = list(DVs_all) + [descriptives_groups] if descriptives_groups is not None else DVs_all
            df, rn = self.data[self.model_names[0]][desc_vars].dropna(), descriptives_groups_rename
            if rn is not None:
                if type(df[descriptives_groups].dropna().iloc[0]) not in [int, str]:
                    df.loc[:, descriptives_groups] = df.loc[:, descriptives_groups].astype(float)
                    if type(list(rn.keys())[0]) is not float:
                        rn = dict([(float(i), rn[i]) for i in rn])
                    df.loc[:, descriptives_groups] = df[descriptives_groups].replace(rn)  # new group labels
                else:
                    rn = dict([(str(int(i)), rn[i]) for i in rn])  # make string (after ensuring int)
                    df.loc[:, descriptives_groups] = df[descriptives_groups].astype(str)  # string in df
                df.loc[:, descriptives_groups] = df[descriptives_groups].replace(rn)  # new group labels
            ds = descriptives_tables(df, self.DV_subsets_dictionary, group=self.descriptives_groups)
            bins = wh(self.binary, True, list(self.DV_subsets_dictionary.keys()))
            conts = wh(self.binary, False, list(self.DV_subsets_dictionary.keys()))
            if type(bins) == list:
                bins = functools.reduce(lambda x, y: x + ' and ' + y, bins)
            if type(conts) == list:
                conts = functools.reduce(lambda x, y: x + ' and ' + y, conts)
            ds = dict(zip([bins, conts], ds))
            tmp = dict(zip([bins, conts], [ds[i].reset_index().assign(i2=i) for i in ds]))
            tmp = [tmp[t].set_index(['i2', 'index']) if 'index' in tmp[t].columns else ds[t] for t in tmp]
            ix = pd.MultiIndex.from_tuples(
                functools.reduce(lambda i, j: i + j, [[('', '')] + list(t.index) for t in tmp]))
            tmp = [t.rename_axis([''] * 2).T.rename_axis([''] * 2).reset_index(1).rename_axis('').T for t in tmp]
            tmp = pd.concat(tmp, ignore_index=True)
            tmp = tmp.set_index(ix)
            tmp.columns = pd.MultiIndex.from_tuples([tuple(i) for i in zip(tmp.columns, tmp.iloc[0, :])])
            tmp = tmp.iloc[1:, :]
            self.descriptives_outcomes = tmp
            capt = string_list(list(self.DV_subsets_dictionary.keys())) + self.caption_suffix
            self.captions.update({'descriptives_outcomes': 'Descriptive Statistics for %s' % capt})
        except Exception as e:
            print(e, 'Could not create descriptives tables.')

    def analyze_correlations(self, digits=3, alpha=0.01):
        """Store & plot results of factor correlatnios with external criteria."""
        self.results_correlations = dict()
        self.results_correlations_p = dict()
        self.results_correlations_tables = dict()
        self.results_correlations_cis = dict()
        pre = 'Correlations between %s Factors and ' % self.analysis
        fsc = [['GF'] + f if self.analysis.upper() == 'BIFACTOR' else f for f in self.fac_series]
        facs = [self.data[self.model_names[f]][fsc[f]] for f in range(len(self.fac_series))]
        self.captions.update({'results_correlations': {}})
        self.captions.update({'results_correlations_p': {}})
        self.captions.update({'results_correlations_tables': {}})
        for j, k in enumerate(self.DV_subsets_dictionary):
            yvs = pd.Series([q if q in self.data[list(self.data.keys())[0]].columns else np.nan
                             for q in self.DV_subsets_dictionary[k]]).dropna()
            if len(yvs) != 0:
                dvs = [self.data[m][yvs] for m in self.model_names]
                out = [corr_dfs(x[0], x[1], digits=digits, alpha=alpha) for x in zip(facs, dvs)]
                out = [pd.concat([o[x] for o in out], keys=self.model_names) for x in [0, 1, 2]]
                self.results_correlations.update({k: out[0]})
                self.results_correlations_p.update({k: out[1]})
                self.results_correlations_tables.update({k: out[2]})
                self.results_correlations_cis.update({k: out[2].applymap(
                    lambda c: [float(r) for r in re.sub('[]]', '', re.sub('.*[[]', '', str(c))).split(', ')])})
                self.captions['results_correlations'].update({k: pre + '%s Outcomes%s' % (k, self.caption_suffix)})
                self.captions['results_correlations_p'].update({k: pre + '%s Outcomes%s' % (k, self.caption_suffix)})
                self.captions['results_correlations_tables'].update({k: pre + '%s Outcomes%s' % (k, self.caption_suffix)})
        try:
            self.results_correlations_cis_bound = pd.concat([self.results_correlations_cis[k] for k in self.results_correlations_cis],
                                                            keys=self.results_correlations_cis.keys()).stack().rename_axis(
                                                                ['Subset', 'Model', 'Factor', 'Outcome']).apply(
                                                                    lambda x: pd.Series(x, index=['LB', 'UB']).rename_axis(
                                                                        'Bound')).stack().to_frame('Correlation')
            self.results_correlations_cis_bound = self.results_correlations_cis_bound.rename(
                dict(zip(self.results_correlations_cis_bound.reset_index().Outcome.unique(),
                         [re.sub('[(].*', '', c) for c in self.results_correlations_cis_bound.reset_index().Outcome.unique()])),
                level='Outcome')
        except Exception as err:
            warnings.warn(f'{err}/n/nCould not create results_correlations_cis_bound attribute.')

    def analyze_regressions(self, alpha=0.01, digits=2, ci_new_line=False):
        """Store & plot results of regressions of factors on external criteria."""
        DVs = [self.DV_subsets_dictionary[k] for k in self.DV_subsets_dictionary.keys()]  # retrieve variable names
        self.alpha = alpha
        # self.results_r2 = [np.nan for i in range(len(DVs))] # for r-squared
        sets = list(self.DV_subsets_dictionary.keys())  # keys for DV subsets
        cic = ['[%s' % '{:.3f}'.format(alpha / 2), '%s]' % '{:.3f}'.format(1 - alpha / 2)]  # CI column names
        self.results_regressions = dict()  # for coefficients
        self.results_regressions_p = dict()  # for bs with *s
        self.results_regressions_objects = dict()  # for fit objects
        self.results_regressions_tables = dict()  # for detailed results tables
        self.results_regressions_cis = dict()  # for 'Estimate [CI_LB-CI_UB]' tables
        results_r2 = []
        fs = [['GF'] + f if self.analysis.upper() == 'BIFACTOR' else f for f in self.fac_series]
        fsc = dict(zip(self.model_names, fs))

        def ci_check(x, cic):  # check that estimate falls with CI
            return (x['Estimate'] > x[cic[0]]) and (x['Estimate'] < x[cic[1]])
        for y in range(len(DVs)):
            kws = {'log': self.binary[y], 'alpha': alpha, 'digits': digits}
            key = list(self.DV_subsets_dictionary.keys())[y]
            out = [try_regress(self.data[m][DVs[y]], self.data[m][fsc[m]], **kws) for m in self.model_names]
            self.results_regressions_objects.update({key: dict(zip(self.model_names, [o[0] for o in out]))})  # out
            # self.results_r2[y] = pd.concat([o[1] for o in out], keys=self.model_names) # r^2
            self.results_regressions.update({sets[y]: pd.concat([o[2] for o in out], keys=self.model_names)})  # bs
            self.results_regressions_p.update({sets[y]: pd.concat([o[3] for o in out], keys=self.model_names)})  # *s
            tabs = pd.concat([o[5] for o in out], keys=self.model_names).reorder_levels([1, 0, 2])  # full tables
            cis = pd.concat([o[6] for o in out], keys=self.model_names).reorder_levels([1, 0, 2]).unstack(0)  # CIs
            ch = tabs.apply(lambda i: ci_check(i, cic), axis=1)  # check estimates are within CIs
            if all(ch) is False:
                warnings.warn('Not all estimates fell within CIs.')
                # warnings.warn('Not all estimates fell within CIs (%s). Check output.' % str(wh(ch, 0, DVs[y])))
                cis = ['Not all estimates fell within CIs (%s). Check output.' % str(wh(ch, 0, DVs))]
            self.results_regressions_tables.update({sets[y]: tabs})
            self.results_regressions_cis.update({sets[y]: cis})
            caps = dict(zip(sets, ['%s Regression Results for %s Outcomes' % (self.analysis, k) for k in sets]))
            ci = ' (Estimate %s[%d%% CI])' % (['', '\n'][ci_new_line], int((1 - alpha) * 100))
            di = dict(zip(['results_regressions%s' % i for i in ['', '_p', '_tables', '_cis']],
                          [caps, caps, caps, dict(zip(sets, [caps[k] + ci for k in caps]))]))
            self.captions.update(**di)  # update dictionary of captions for table attributes
            rsq = tabs[wh(['R-Squared' in c for c in tabs.columns], 1, list(tabs.columns))]
            results_r2 = results_r2 + [rsq.replace('', np.nan).dropna()]
        self.results_r2 = pd.concat(results_r2, axis=0).reset_index(2, drop=True).unstack()


# %% Documentation Class

class Factor_Model_Results(Factor_Model_Analysis):
    r"""
    Results.

    Examples
    --------
    >>> # Re-Write MPlus Syntax?
    >>> write_syntax = True

    >>> # Re-Run MPlus:
    >>> cfas_run = False
    >>> cfas_all_run = False
    >>> cfas_PTSD_run = False
    >>> bfs_run = False

    >>> # Validation Analyses
    >>> DVs_sf = ['Calm', 'Depressed', 'Less_Accomplished', 'Less_Careful']
    >>> DVs_dx = ['PTSD', 'MDD', 'GAD', 'Phobia', 'Panic', 'AUD', 'NUD', 'BPD', 'CD', 'ASPD']
    >>> DVs_drink = ['MAXDRINKS', 'DRINKFREQ', 'USUALAMT', 'INTOX', 'FREQMAX', 'BINGE']
    >>> DV_subsets_dictionary = {'SF-12': DVs_sf, 'Diagnoses': DVs_dx, 'Drinking': DVs_drink} # dictionary of DV subsets
    >>> logistic_regression = [False, True, False] # logistic regression for 2nd DV subset

    >>> # Documentation
    >>> file_source = '/media/ena/My Passport/Manuscripts/manuscript_PTSD_revision.docx' # base file for revision
    >>> file_source_supp = '/media/ena/My Passport/Manuscripts/manuscript_PTSD_supplement.docx' # same for supplement
    >>> new_doc_file = 'manuscript_PTSD_revised.docx' # where to write modifications
    >>> new_doc_file_supp = 'manuscript_PTSD_revised.docx' # where to write modifications

    >>> # Arguments:
    >>> args_cfas = {'model_dictionary': model_dictionary, 'model_names': model_names, 'fac_series': fac_series,
    ...             'data_file': data_file, 'categorical': True, 'TYPE': 'complex',
    ...             'indicator_dictionary': indicator_dict, 'free_load': False,
    ...             'cluster': 'VARUNIT', 'stratification': 'VARSTRAT', 'weight': 'AUDWEIGHT',
    ...             'run_mplus_syntax': cfas_run, 'DV_subsets_dictionary': DV_subsets_dictionary, 'base_model': 0,
    ...             'digits_r': 3, 'descriptives_groups': 'Sex', 'digits_d': 2}
    >>> args_cfas_all = {'model_dictionary': model_dictionary_all, 'model_names': model_names, 'fac_series': fac_series,
    ...                 'data_file': data_file, 'categorical': True, 'indicator_dictionary': indicator_dict,
    ...                 'weight': 'AUDWEIGHT', 'TYPE': 'complex', 'free_load': False,
    ...                 'cluster': 'VARUNIT', 'stratification': 'VARSTRAT',
    ...                 'run_mplus_syntax': cfas_all_run, 'DV_subsets_dictionary': DV_subsets_dictionary,
    ...                 'base_model': 0, 'digits_r': 3, 'descriptives_groups': 'Sex',
    ...                'digits_d': 2, 'write_syntax': write_syntax}
    >>> args_cfas_PTSD = {'model_dictionary': model_dictionary, 'model_names': model_names, 'data_file': data_file,
    ...                  'fac_series': fac_series, data_file: data_file_PTSD, 'estimator': 'mlr', # LPTSD only
    ...                  'file_suffix': 'PTSD', 'categorical': True, 'indicator_dictionary': None, 'free_load': False,
    ...                  'run': cfas_PTSD_run, 'DV_subsets_dictionary': DV_subsets_dictionary,
    ...                 'write_syntax': write_syntax,
    ...                  'base_model': 0, 'digits_r': 3, 'descriptives_groups': 'Sex', 'digits_d': 2}
    ...                  # subpopulation: 'PTSD == 1' # using subsetted data instead
    >>> args_bf = {'model_dictionary': model_dictionary, 'model_names': model_names, 'data_file': data_file,
    ...           'fac_series': fac_series, 'categorical': True, 'analysis': 'Bifactor', 'run_mplus_syntax': bfs_run,
    ...           'indicator_dictionary': indicator_dict, 'free_load': False, 'estimator': 'mlr', 'cores': 4,
    ...           'TYPE': 'complex', 'cluster': 'VARUNIT', 'stratification': 'VARSTRAT', 'weight': 'AUDWEIGHT',
    ...           'analysis_statement_extras':  'MCONVERGENCE: .1; \n\t LOGCRITERION: .02; \n\t RLOGCRITERION: .0001;',
    ...          'DV_subsets_dictionary': DV_subsets_dictionary, 'base_model': 0, 'digits_r': 3,
    ...           'descriptives_groups': 'Sex', 'digits_d': 2, 'write_syntax': write_syntax,}
    >>> args_documentation = {'analyses_types': ['correlations', 'regressions'], 'binary': [False, True, False],
    ...                      'source_file': file_source,
    ...                      'supplement_include': True, 'source_file_supplement': file_source_supp,
    ...                      'output_file': new_doc_file, 'output_file_supplement': new_doc_file_supp,
    ...                      'digits_d': 2, 'rename_dict': None, 'abbreviation_dict': None} # documentation arguments
    >>> args_cfas.update(args_documentation)
    >>> args_cfas_all.update(args_documentation)
    >>> args_cfas_PTSD.update(args_documentation)
    >>> args_bf.update(args_documentation)

    >>> # Make Objects
    >>> cfas = Factor_Model_Results(**args_cfas)
    >>> cfas_all = Factor_Model_Results(**args_cfas_all)
    >>> cfas_PTSD = Factor_Model_Results(**args_cfas_PTSD)
    >>> bfs = Factor_Model_Results(**args_bf)

    """

    def __init__(self, alpha=0.01, rename_dict=None, abbreviation_dict=None, hypotheses_dict=None,
                 analyses_types=['correlations', 'regressions'], gather_file_attributes=True,
                 descriptives_groups=None, descriptives_groups_rename=None, DVs_all=None, **kwargs):
        """Initialize (passes on arguments & keyword arguments to parent class Factor_Model_Analysis)."""
        super(Factor_Model_Results, self).__init__(**kwargs)  # initialize parent class object
        self.analyses_types = analyses_types
        self.alpha = alpha
        self.dictionaries = {'rename': rename_dict, 'abbreviations': abbreviation_dict, 'hypotheses': hypotheses_dict}
        self.analyze_descriptives(DVs_all=DVs_all, descriptives_groups=descriptives_groups,
                                  descriptives_groups_rename=descriptives_groups_rename)  # descriptive statistics
        from functions_documentation import string_list
        if 'descriptives_outcomes' not in self.captions.keys():
            capt = string_list(list(self.DV_subsets_dictionary.keys())) + self.caption_suffix
            self.captions.update({'descriptives_outcomes': 'Descriptive Statistics for %s' % capt})
        for a in self.analyses_types:
            try:
                if a == 'regressions':
                    self.analyze_regressions(alpha=self.alpha)
                if a == 'correlations':
                    self.analyze_correlations(alpha=self.alpha)
            except Exception as e:
                print(e, '\n Analyzing %s failed.' % a)

    def initialize_documents(self, directory='Manuscripts', source_files=None,
                             sections=None, title='Introduction',
                             # ['Title', 'Abstract', 'Introduction', 'Methods', 'Results',
                             #  'Discussion', 'References', 'Tables', 'Figures']
                             source_files_supplement=None, sections_supplement=None,
                             output_sections=True,
                             title_page_header=None,
                             output_sections_supp=True,
                             # ['Methods', 'Results', 'Tables']
                             header_titles=True, font_name='Times New Roman', font_size=12):
        """Initialize manuscript (& supplement if desired)."""
        if output_sections:
            output_sections = copy.deepcopy(sections)
        if output_sections_supp:
            output_sections_supp = copy.deepcopy(sections_supplement)
        if type(header_titles) is bool:
            if header_titles:
                secs = list([copy.deepcopy(output_sections), copy.deepcopy(output_sections_supp)])
                hs = [copy.deepcopy(x) for x in secs]  # default heads named ~ sections
                header_titles = copy.deepcopy(hs)
        for i, h_tit in enumerate(header_titles):
            print(output_sections)
            if ('introduction' in [h.lower() for h in h_tit]) and (title is not None):  # introduction head = title
                header_titles[i][wh([h.lower() for h in h_tit], 'introduction')] = title
            if ('title' in [h.lower() for h in h_tit]):  # title seciton header = title
                header_titles[i][wh([h.lower() for h in h_tit], 'title')] = ''  # no title page head
        if 'Manuscripts' not in os.listdir():
            os.system('mkdir Manuscripts')
        self.main = {}
        self.supplement = {}
        self.documents = {'main': {}} if source_files_supplement is None else {'main': {}, 'supplement': {}}
        if directory is None:
            directory = os.getcwd()
        for x in range(len(output_sections)):  # iterate sections
            new_sec_file = '%s/%s.docx' % (directory, output_sections[x])
            print('\nInitializing main section %s...\nWriting to %s...' % (output_sections[x], new_sec_file))
            if output_sections[x] not in sections:
                new_sec = docx.Document()  # start document
            else:
                loc = wh(sections, output_sections[x], source_files)  # section source file
                new_sec = docx.Document(loc)  # start with source
            if header_titles[0] is not None:
                # if len(new_sec.paragraphs) == 0: new_sec.add_paragraph()
                # sec_par = new_sec.paragraphs[0].insert_paragraph_before()
                if len(new_sec.paragraphs) == 0:
                    sec_par = new_sec.add_paragraph()
                else:
                    sec_par = new_sec.paragraphs[0].insert_paragraph_before()
                sec_par.line_spacing_rule = WD_LINE_SPACING.SINGLE
                sec_head = sec_par.add_run(header_titles[0][x])
                sec_head.bold = True
                # new_sec.header.text = header_titles[0][x]
                # new_sec.header.is_linked_to_previous = True
                # for p in : # iterate paragraphs in section
                sec_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sec_par.paragraph_format.space_before = False
                sec_par.paragraph_format.space_after = False
                for r in sec_par.runs:  # iterate runs within section paragraph
                    r.font.name = font_name
                    r.font.size = font_size
                    r.line_spacing_rule = WD_LINE_SPACING.SINGLE
            self.main.update({output_sections[x]: new_sec})  # add to dictionary of section documents
            self.documents['main'].update({output_sections[x]: new_sec_file})  # add section file name
            self.main[output_sections[x]].save(new_sec_file)
        self.supplement = {}
        if source_files_supplement is not None:
            if 'Supplement' not in os.listdir(directory):
                os.system('mkdir %s/Supplement' % directory)
            for x in range(len(output_sections_supp)):  # iterate sections
                new_sec_file = 'Manuscripts/Supplement/%s.docx' % (output_sections_supp[x])
                print('\nInitializing supplement section %s...' % output_sections_supp[x])
                print('Writing to %s...' % new_sec_file)
                if output_sections_supp[x] not in sections_supplement:
                    new_sec = docx.Document()  # start document
                else:
                    loc = wh(sections_supplement, output_sections_supp[x], source_files_supplement)  # section source
                    new_sec = docx.Document(loc)  # start with source
                if header_titles[1] is not None:
                    # if len(new_sec.paragraphs) == 0: new_sec.add_paragraph()
                    # sec_par = new_sec.paragraphs[0].insert_paragraph_before()
                    if len(new_sec.paragraphs) == 0:
                        sec_par = new_sec.add_paragraph()
                    else:
                        sec_par = new_sec.paragraphs[0].insert_paragraph_before()
                    sec_par.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    sec_head = sec_par.add_run(header_titles[1][x])
                    # new_sec.header.text = header_titles[0][x]
                    # new_sec.header.is_linked_to_previous = True
                    # for p in : # iterate paragraphs in section
                    sec_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sec_par.paragraph_format.space_before = False
                    sec_par.paragraph_format.space_after = False
                    for r in sec_par.runs:  # iterate runs within section paragraph
                        r.font.name = font_name
                        r.font.size = font_size
                        r.line_spacing_rule = WD_LINE_SPACING.SINGLE
                self.supplement.update({output_sections_supp[x]: new_sec})  # add to dictionary of section documents
                self.documents['supplement'].update({output_sections_supp[x]: new_sec_file})  # add section file name
                self.supplement[output_sections_supp[x]].save(new_sec_file)
        self.table_start, self.table_start_supplement = 1, 1  # assume start at table 1
        documentation = [dict_part(self.main, 'items')]
        if len(self.supplement) > 0:
            documentation + [dict_part(self.supplement, 'items')]
        for d in range(len(documentation)):
            for dc in documentation[d]:
                for i in range(len(dc.paragraphs)):  # find 1st table in source document to determine table start number
                    if '.*Table S?([0-9]*).\n.*' not in dc.paragraphs[-i].text:
                        continue
                    else:
                        start = re.sub(r'\n', '', re.sub('.*Table S?([0-9]*).\n.*', '\\1', dc.paragraphs[-i].text))
                        if d == 0:
                            self.table_start = float(start)
                        else:
                            self.table_start_supplement = float(start)
                        break

    def define_styles(self, font_name='Times New Roman', style_type=None, font_size=12,
                      parts=['main', 'supplement']):
        """Define document styles."""
        if type(parts) is str:
            parts = [parts]
        for i, s in enumerate(parts):
            for k in [self.main, self.supplement][i]:
                for p in [self.main, self.supplement][i][k].paragraphs:
                    for r in p.runs:
                        r.font.name = font_name
                        r.font.size = Pt(font_size)
                [self.main, self.supplement][i][k].save(self.documents[s][k])

    def compare_objects(self, comparison_model, attributes, object_names=['Self', 'Other'],
                        keep_dims=True, keep_same=False, sig='**',
                        store_as_attribute=True, clear_comparisons_attribute=False, plot=True,
                        font={'family': 'serif', 'size': 10}, title_fontsize=12,
                        cb_map='coolwarm', cbar_min_range=[-0.3, 0.3],
                        show=True, save=True, p_annotate=True):
        """Compare different objects of factor results class."""
        if type(p_annotate) == bool:
            p_annotate = [p_annotate] * len(attributes)  # if 1 specified, assume for all
        if type(plot) == bool:
            plot = [plot] * len(attributes)  # if 1 specified, assume for all
        if clear_comparisons_attribute:  # start over comparisons attribute if desired (clear results of past runs)
            if 'results_compare_analyses' in dir(self):
                self.__delattr__('results_compare_analyses')
        if type(attributes) == str:
            attributes = [attributes]  # convert to list if only specified 1 attribute
        init = ('results_compare_analyses' not in dir(self)) or (self.results_compare_analyses is None)
        if init:  # start dictionary of lists of dictionaries attribute if absent (see above)
            self.results_compare_analyses = {object_names[1]: [{}, {}, {}]}
        elif object_names[1] not in list(self.results_compare_analyses.keys()):  # if comparison model not in keys, add
            self.results_compare_analyses.update({object_names[1]: [{}, {}, {}]})
        fig_files = []
        if 'Plots/Differences' not in os.listdir():
            os.system('mkdir Plots/Differences')
        for att in attributes:
            df_1, df_2 = self.__getattribute__(att), comparison_model.__getattribute__(att)  # main & comparison
            df_1, df_2 = [dict_part(x, 'items') if type(x) == dict else x for x in [df_1, df_2]]  # items if dictionary
            out = compare_data(df_1, df_2, keep_dims=keep_dims, keep_same=keep_same, object_names=object_names)
            if store_as_attribute:  # store as attributes if desired
                [self.results_compare_analyses[object_names[1]][t].update({att: out[t]}) for t in range(len(out))]
            if (plot[wh(attributes, att)]) and all([d.empty is False for d in out[1]]):
                p_ann = p_annotate[wh(attributes, att)]
                if ('%s_p' % att not in dir(self)) and (p_ann):
                    print('%s_p unavailable' % att)
                    continue
                ant = self.__getattribute__(att + '_p') if p_ann else True  # annotate = True or df with N.S.=0
                bnt = comparison_model.__getattribute__(att + '_p') if p_ann else True  # comparison p-stars
                if isinstance(ant, dict):
                    ant = dict_part(ant, 'items')  # get dfs from dictionary if needed
                file_stem = 'Plots/Differences/%s_%s_%s' % (att, object_names[0], object_names[1])
                vmin = min([min(d.min()) for d in out[1]] + [cbar_min_range[0]])  # color bar maximum
                vmax = max([max(d.max()) for d in out[1]] + [cbar_min_range[1]])  # color bar minimum
                with PdfPages(file_stem + '.pdf') as pdf_compare:  # "gathering" figures to save
                    for i in range(len(out[1])):
                        an = out[1][i]  # annotate with original differences
                        if p_ann:
                            x = list(bnt.keys())[i]
                            non_sig_a = ant[i].applymap(lambda d: 1 if str(d).count('*') < sig.count('*') else 0)
                            non_sig_b = bnt[x].applymap(lambda d: 1 if str(d).count('*') < sig.count('*') else 0)
                            ix = non_sig_b.index.intersection(non_sig_a.index)
                            cols = non_sig_b.columns.intersection(non_sig_a.columns)
                            non_sig = non_sig_a.add(non_sig_b.loc[ix, cols])  # add N.S. 0/1
                            non_sig = non_sig.replace(2, np.nan).replace(1, 0)  # don't change data unless both N.S.
                            df = an.add(non_sig)  # N.S. -> NaN (because N.S. -> NaN above & rest stay same b/c + 0)
                        if ('regression' in att.lower()) and any(df.index.to_frame().iloc[:, 1] == 'b0'):
                            df = df.drop(labels='b0', axis=0, level=1)  # remove intercept rows
                            an = an.drop(labels='b0', axis=0, level=1)  # ""
                        if (an.shape[0] != df.shape[0]) or (an.shape[1] != df.shape[1]):
                            an = True
                            warnings.warn('Could not use p-annotation for %s. Shapes do not match.' % att)
                        fig = heatmaps(df, self.model_names, data_annotate=an, cb_map=cb_map, colorbar_center=0,
                                       show=show, save=save, col_rename_dict=self.dictionaries['rename'],
                                       font=font, title_fontsize=title_fontsize, cbar_min_range=[vmin, vmax])
                        if save:
                            orient = 'landscape' if len(df.columns) > 6 else 'portrait'
                            pdf_compare.savefig(orientation=orient)  # save in multi-page PDF
                            fig_files = fig_files + [file_stem + '.pdf']
                if show:
                    plt.get_current_fig_manager().window.showMaximized()  # maximize plot window
                    fig.show()
                else:
                    plt.close('all')
                if 'files' not in dir(self):
                    self.files = {}
        if plot:
            self.files['figures'].update({'results_comparisons': fig_files})  # add file names to dictionary

    # def plot_objects_correlations(self, comparison_model_list, mask_NS=True, model_names=['Self', 'A', 'B'],
    #                               save=True, show=True):
    #     """Plot Correlation Comparisons."""
    #     mods = [self] + comparison_model_list
    #     if 'Plots/Comparisons' not in os.listdir():
    #         os.system('mkdir Plots/Comparisons')
    #     stem = 'Plots/Comparisons/results_correlations_' + functools.reduce(
    #         lambda x, y: '%s_%s' % (x, y), model_names)
    #     with pdfPages(stem + '.pdf') as pdf:
    #         for subset in list(self.DV_subsets_dictionary.keys()):
    #             r = [corr_dfs(mods[m].df_factor_scores, mods[m].data[subset])[1] for m in range(len(mods))]
    #             j = [mat.loc[x.index.Intersection(r[0].index), x.columns.Intersection(r[0].columns)] for x in r[1:]]
    #             corrs = [r[0]] + j
    #             fig, axes = plt.subplots(nrows=1, ncols=len(mods))
    #             for m in range(len(mods)):
    #                 da = corrs_sf[m].apply(lambda x: x.str.strip('*').astype(float))
    #                 df = df.apply(lambda x: x.str.strip('*').astype(float))
    #                 if mask_NS:
    #                     df = corrs_sf[m].applymap(lambda x: np.nan if x.count('*') < sig.count('*') else x)
    #                 sb.heatmap(df, annot=da, ax=axes[m], cbar=True, cmap='coolwarm')
    #                 axes[m].tick_params(top=True, bottom=False, left=True if m == 0 else False,
    #                                     labeltop=True, labelbottom=False)
    #                 plt.setp(axes[m].get_xticklabels(), rotation=-30, ha='right',
    #                          rotation_mode='anchor', fontstyle='italic')
    #                 axes[m].tick_params(which='minor', bottom=False, left=True if m == 0 else False)
    #                 axes[m].grid(which='minor', color='w', linestyle='-', linewidth=3)
    #         if save:
    #             pdf.savefig(stem + '.pdf')
    #             self.files['figures'].update({stem + '.pdf'})

    def plot_correlation_cis(self, col_wrap=4, file_stem=None, plot_dir='Plots/'):
        """Plot correlation boxplots based on confidence intervals."""
        if 'figures' not in dir(self):
            self.figures = {}
        args = dict(y='Correlation', x='Factor', hue='Outcome', col='Model', kind='box', sharex=False, col_wrap=col_wrap)
        figs = [sb.catplot(data=self.results_correlations_cis_bound.loc[k].reset_index(), **args) for k in self.DV_subsets_dictionary]
        self.figures.update({'Correlation Confidence Intervals': dict(zip(self.DV_subsets_dictionary.keys(), figs))})
        if file_stem is not None:
            [self.figures['Correlation Confidence Intervals'][k].savefig(f'{plot_dir}{file_stem}_{k}.jpeg')
             for k in self.figures['Correlation Confidence Intervals']]

    def plot_heatmaps(self, plot_what, attribute_key=None, comparison_model=None, p_annotate=True,
                      cb_map='coolwarm', cbar_min_range=[-0.3, 0.3], figure_location='main',
                      show=True, save=True, file_extension='.pdf', main=False, supplement=True,
                      fig_title=None, sig='**', font={'family': 'serif', 'size': 10},
                      title_fontsize=12, save_individual=True, directory=None, plot_fac_cors=True):
        """
        Plot heat maps of data in attribute.

        Parameters
        ----------
        plot_what : str
            Name of attribute to visualize with heat maps.
        directory : str, optional
            Sub-directory (including final /) under plots to which to save figure. Default, top-level directory.
        attribute_key : str, optional
            If the attribute listed in plot_what is a dictionary, specify the key to access the data to plot.
        comparison_model : Factor_Model_Results object, optional
            Model for attributes comparison. The default is None.
        p_annotate : bool, optional
            Whether or not to include p-value stars in heat map cell annotations. The default is False.
        cb_map : str, optional
            Color bar color map. The default is 'coolwarm'.
        cbar_min_range : list, optional
            Minimum range mapped by colorbar (will be extended to accommodate data if needed). Default = [-0.3, 0.3].
        show : bool, optional
            Whether to display figures. The default is True.
        file_extension : str, optional
            File extension/format for saved figure (e.g. '.pdf', '.eps', '.jpeg). The default is '.pdf'.
        figure_location : str, optional
            Key in self.files dictionary in which to store location of saved figures (default is 'main').
            Specify 'main' or 'supplement' to flag for later documentation in main manuscript or supplement.
        supplement : bool, optional
            Whether or not to number figure & store it as supplement figure (rather than main). The default is True.
        fig_title : str, optional
             Specify to override default figure title stem (default based just on analysis type). The default is None.
        sig : str, optional
            Minimum of stars indicating significance. The default is '**'.
        font : str, optional
            Font argument dictionary. The default is {'family': 'serif', 'size': 10}.
        title_fontsize : str, optional
            Font size for figure/subplot titles. The default is 12.
        save : bool, optional
            Save figures as files?. The default is True.
        save_individual : bool, optional
            Set as True to save individual subset .eps files (along with multi-page PDF). The default is False.
        plot_fac_cors : bool, optional
            Whether to include a subplot with distributions of factor correlations. Default is True.

        """
        what = plot_what.lower() + '_p' if p_annotate else plot_what.lower()
        if directory is None:
            directory = what + '/'
        if 'Plots' not in os.listdir():
            os.system('mkdir Plots')
        if directory[:-1] not in os.listdir('Plots'):
            os.system('mkdir Plots/%s' % directory)
        s1 = re.sub('_0', '', '%s_%s' % (self.analysis, self.files['file_suffix']))  # suffix
        # fac_num_frac = [len(f) / sum([len(f) for f in self.fac_series]) for f in self.fac_series]  # s of factors/total
        # fac_frac = [f / max(fac_num_frac) for f in fac_num_frac]  # proportion relative to model with most factors
        dfs = self.__getattribute__(what)  # main list of data to plot
        if attribute_key is not None:
            dfs = dfs[attribute_key]  # retrieve by key if needed
        if type(dfs) == dict:
            dfs = [a for a in dict_part(dfs, 'items')]
        if type(dfs) is not list:
            dfs = [dfs]
        file_stem = '%s/Plots/%s%s_%s' % (os.getcwd(), directory, plot_what, s1)  # file
        if p_annotate:  # raw data for annotations; N.S. values -> 0 for color bar (if desired)
            dann = [d.apply(lambda x: x.str.strip('*'), axis=1).astype(float) for d in dfs]
            dfs = [d.applymap(lambda x: x.strip('*') if str(x).count('*') >= sig.count('*') else np.nan) for d in dfs]
        data_list = [d.astype(float) for d in dfs]
        fig_files = []  # to hold figure file names
        if plot_what.lower() == 'descriptives_factors':
            fig = square_heatmap(data_list, self.model_names)
            fig.savefig(file_stem + file_extension)
            fig_files = fig_files + [file_stem + file_extension]
        else:
            with PdfPages(file_stem + '.pdf') as pdf_compare:  # "gathering" figures to save
                for i in range(len(data_list)):
                    df_ann = dann[i] if p_annotate else True
                    df = data_list[i]
                    subset = list(self.__getattribute__(what).keys())[i]  # subset name
                    bin_dict = dict(zip(list(self.DV_subsets_dictionary.keys()), self.binary))  # binary outcome?
                    center = 1 if bin_dict[subset] and ('regression' in plot_what.lower()) else 0
                    if ('regression' in plot_what.lower()) and any(df.index.to_frame().iloc[:, 1] == 'b0'):  # b0
                        df = df.drop(labels='b0', axis=0, level=1)  # remove intercept rows
                        df_ann = df_ann.drop(labels='b0', axis=0, level=1)  # ""
                    if fig_title is None:
                        fig_title = cap(s1) + ' ' + cap(plot_what)
                    fig_head = None if fig_title is None else '%s for %s' % (fig_title, subset)
                    fig = heatmaps(df, self.model_names, data_annotate=df_ann, fig_title=fig_head, cb_map=cb_map,
                                   colorbar_center=center,  # see above code -- center colorbar on 1 if logistic, else 0
                                   show=show, save=save, col_rename_dict=self.dictionaries['rename'],
                                   font=font, title_fontsize=title_fontsize,
                                   # intercors_label='Factor',
                                   # intercors=[self.data[m][self.fac_series[m]] for m in self.model_names],
                                   intercors_label='Distributions of',
                                   plot_ranges=False)
                    if save:
                        orient = 'landscape' if len(df.columns) > 6 else 'portrait'
                        if save_individual:  # if want to save individual subset files (along with multi-page PDF)
                            fig.savefig('%s_%s.jpg' % (file_stem, subset))
                            way = 'portrait' if orient == 'landscape' else 'landscape'
                            fig_files = fig_files + [('%s_%s.jpg' % (file_stem, subset), way)]  # store figure file name
                        pdf_compare.savefig(orientation=orient)  # save in multi-page PDF
            if show:
                plt.get_current_fig_manager().window.showMaximized()  # maximize plot window
                fig.show()
            else:
                plt.close('all')
            if 'figures' not in self.files.keys():
                self.files.update({'figures': {}})
            if figure_location not in list(self.files['figures'].keys()):
                self.files['figures'].update({figure_location: {}})
            self.files['figures'][figure_location].update({plot_what: fig_files})

    def plot_heatmaps_contrast(self, factor_contrasts, correlates, super_title=True, save=True,
                               colorbar=True, cb_map='coolwarm', cb_center=0, cbar_min_range=[-0.3, 0.3],
                               cbar_kwargs={'orientation': 'vertical', 'shrink': 1,
                                            'extend': 'min', 'extendfrac': 0.1, 'drawedges': False}):
        """Heat maps contrasting factor correlations."""
        suff = '_' + self.files['file_suffix'] if self.files['file_suffix'] != 0 else ''  # file suffix
        if 'Contrasts' not in os.listdir('Plots'):
            os.system('mkdir Plots/Contrasts')
        file = '%s/Plots/Contrasts/%s_%s%s.pdf' % (os.getcwd(), 'results_correlations_', self.analysis, suff)
        with PdfPages(file) as pdf:
            for facs in factor_contrasts:  # iterate factor contrasts
                mods = self.fac_series[self.fac_series.apply(lambda x: all([i in x for i in facs]))].index.values
                fig, axes = plt.subplots(**dict(zip(['nrows', 'ncols'], square_grid(len(mods)))), figsize=[25, 15])
                if super_title:
                    fig.suptitle(functools.reduce(lambda x, y: x + ' versus ' + y, facs))
                for m in range(len(mods)):  # iterate models with contrasting factors
                    ax = axes.ravel()[m]  # mth axes out of flattened axis array
                    DVs = [correlates[f] if f in correlates.index.values else None for f in facs]  # DVs of interest
                    DVs = pd.unique(pd.Series(DVs).explode().dropna())  # flatten list
                    if None in DVs:
                        DVs.remove(None)
                    df = self.data[self.model_names[m]]  # factor scores
                    corr_mat_dat = df.corr()  # correlation matrix
                    ixs = list(corr_mat_dat.index.values)  # list of all indices (DVs & model-factors)
                    [ixs.remove(i) for i in DVs]  # leave just DV tuple index values in list
                    corr_m = corr_mat_dat.loc[DVs, ixs]  # correlations with DVs in rows, model-factors in columns
                    corr_m = corr_m.rename(self.dictionaries['rename'], axis=1).rename(self.dictionaries['rename'])
                    corr_ann = corr_m  # store original corr_m so can annotate with values even if greyed out because N.S.
                    for r in DVs:
                        for c in ixs:
                            drop = np.logical_or(np.array((pd.isnull(df[r]))), np.array(pd.isnull(df[c])))  # NAs
                            if sp.stats.pearsonr(df[~drop][r], df[~drop][c])[0] >= self.alpha:  # if not significant...
                                if (r in corr_m.index.values) and (c in corr_m.columns):
                                    corr_m.loc[r][c] = 0  # replace with 0 to guide color-bar
                    if m == 0:  # for 1st model, sort with highest values at top
                        corr_m = corr_m.sort_values(by=corr_m.columns[0], ascending=False)
                        sorted_ixs = corr_m.index.values  # store index order so can remain consistent across models
                    else:  # subsequent models, order like the 1st
                        corr_m = corr_m.loc[sorted_ixs]
                    corr_ann = corr_ann.loc[sorted_ixs]  # sort annotation version, too
                    # sb.clustermap(corr_m, cbar_pos=None, cmap='coolwarm')
                    cmin, cmax = colorbar_min_range(corr_m, cbar_min_range=cbar_min_range)  # color bar range
                    cbar = colorbar if m == len(mods) - 1 else None  # include colorbar if at last model
                    cb_key = {**cbar_kwargs, 'use_gridspec': True} if m == len(mods) - 1 else None
                    sb.heatmap(corr_m, xticklabels=facs, ax=ax, annot=corr_ann, fmt='.2f',
                               cbar=cbar, cmap=cb_map, center=cb_center, vmin=cmin, vmax=cmax, cbar_kws=cb_key)
                    ax.tick_params(top=True, bottom=False, left=True, labeltop=True, labelbottom=False)  # labels: top
                    plt.setp(ax.get_xticklabels(), rotation=-30, ha='right', rotation_mode='anchor')  # ticks rotate
                    ax.tick_params(which='minor', bottom=False, left=False)
                    ax.grid(which='minor', color='w', linestyle='-', linewidth=3)
                    sub = re.sub('_', ' ', mods[m]).title() + ' ' + self.analysis  # model name & analysis to subtitle
                    if self.files['file_suffix'] != 0:
                        sub = sub + ' (' + self.files['file_suffix'] + ')'  # (suffix)
                    ax.set_title(cap_title([mods[m], self.analysis]))  # subplot title
                    for edge, spine in ax.spines.items():  # turn off spines & edges
                        spine.set_visible(False)
                    ax.grid(b=False)
                for edge, spine in ax.spines.items():  # turn off spines & edges
                    spine.set_visible(False)
                plt.subplots_adjust(left=0.03, bottom=0.12, top=0.85, right=0.97, hspace=0.15, wspace=0.15)  # adjust
                plt.get_current_fig_manager().window.showMaximized()
                fig.tight_layout()
                if save:
                    pdf.savefig()
                fig.show()

    # def plot_kde_3d(self, DVs, model_num, factors=None,
    #                 fig_n_start=None, show=True, store=False, supplement=True):
    #     """Joint kernel density plots (partly from towardsdatascience.com)."""
    #     if fig_n_start is None:
    #         if supplement:
    #             fig_n_start = len(self.figures_supplement) + 1 if 'figures_supplement' in dir(self) else 1  # tracking
    #         else:
    #             fig_n_start = len(self.figures) + 1 if 'figures' in dir(self) else 1  # start tracking figure #
    #     if factors is None:
    #         factors = self.fac_series[model_num]  # plot all model factors if not specified
    #     for f in factors:
    #         fig = plot_kde(self.data.loc[:, self.model_names[model_num]], DVs, f, self.model_names[model_num])
    #     if store:
    #         if supplement:
    #             started = 'figures_supplement' in dir(self)
    #             self.figures_supplement = self.figures_supplement + [fig] if started else [fig]  # store
    #         else:
    #             self.figures = self.figures + [fig] if 'figures' in dir(self) else [fig]  # store attribute
    #     if show:
    #         plt.get_current_fig_manager().window.showMaximized()  # maximize plot window
    #         fig.show()
    #     else:
    #         plt.close('all')

    def plot_boxplots(self, data_dictionary,
                      attribute='results_correlations', att_lab='Correlations', cats_by='Factor',
                      magnitude=True, rename=True, exclude_vars=None,
                      show=True, save=True, figsize=[26, 10], sharey=True,
                      legend_cols=None, legend_loc='lower center', legend_mode='expand', fig_legend=True,
                      ylim=None, p_annotate=True, **kwargs):
        """Plot Box Plots of Correlation or Regression Coefficients."""
        attrib = self.__getattribute__(attribute)
        fig, axes = plt.subplots(*square_grid(len(attrib)), figsize=figsize, sharey=sharey)
        for i, k in enumerate(attrib):
            df = attrib[k]
            if exclude_vars is not None:
                if type(exclude_vars) is str:
                    exclude_vars = [exclude_vars]  # ensure iterable
                for e in exclude_vars:
                    if e in df.columns:
                        df = df.drop(e, axis=1)  # drop certain variables (if desired)
            if (rename) and ('rename' in self.dictionaries.keys()):  # use internal rename (if not provided)
                rename = self.dictionaries['rename']
            else:
                rename = self.rename_dict
            if rename is not None:  # re-name (if desired)
                df = df.rename(rename, axis=0).rename(rename, axis=1)
            if magnitude:
                df = df.abs()
            ix_names = ['Model', 'Factor'] if all([q is None for q in df.index.names]) else list(df.index.names)
            cat_mat = pd.DataFrame(df.stack().rename_axis(ix_names + ['Outcome']), columns=[att_lab])
            sb.boxenplot(data=cat_mat.reset_index(), x=cats_by, y=att_lab, hue=ix_names[0],
                         ax=axes.ravel()[i], trust_alpha=self.alpha, **kwargs)
            axes.ravel()[i].set_ylabel(att_lab + ['', ' (Magnitude)'][magnitude])
            if ylim is not None:
                axes.ravel()[i].set_ylim(ylim)
            axes.ravel()[i].set_title(k)
            if fig_legend:
                axes.ravel()[i].legend_.remove()
            else:
                if legend_cols is None:
                    legend_cols = len(pd.unique(cat_mat.reset_index()[ix_names[0]]))
                axes.ravel()[i].legend(ncol=legend_cols, loc=legend_loc, frameon=False)
        if len(self.model_names) is None:
            legend_cols = len(self.model_names)
        if fig_legend:  # figure-level legend (if desired)
            fig.legend(*[q for q in axes.ravel()[i].get_legend_handles_labels()], loc=legend_loc,
                       mode=legend_mode, ncol=legend_cols, frameon=False)  # legend (horizontal)
        if show:
            plt.get_current_fig_manager().window.showMaximized()  # maximize plot window
            if fig_legend is False:
                fig.tight_layout()
            fig.show()
        else:
            plt.close('all')
        if save:
            caption = 'Boxplots of %s Model-Factors\' %s' % (self.analysis, att_lab)
            if self.files['file_suffix'] != 0:
                caption = caption + ' (' + self.files['file_suffix'] + ')'  # (suffix)
            if 'figures' not in self.files.keys():
                self.files.update({'figures': {}})
            if attribute not in self.captions.keys():
                self.captions.update({attribute: {}})
            if 'Plots' not in os.listdir():
                os.system('mkdir Plots')  # make Plots directory if absent
            suf = '_' + self.files['file_suffix'] if self.files['file_suffix'] != 0 else ''
            file = 'Plots/Boxplots_%s_%s_%s%s.png' % (attribute, cats_by, self.analysis, suf)
            self.files['figures'].update({'Boxplots_%s' % attribute: file})
            if exclude_vars is not None:
                caption = caption + ' (Excluding ' + string_list(exclude_vars) + ')'
            caption = re.sub('[)] [(]', '; ', caption)
            self.captions[attribute].update({'Boxplots': caption})
            fig.savefig(file)

    def plot_boxplots_compare(self, comparison_model, model_keys=['Self', 'Other'], comparison_type='Sample',
                              dv_groups=None, dv_groups_keys=None, fac_groups=None, fac_groups_keys=None,
                              reverse_code=None, attribute='results_correlations', att_lab='Correlations',
                              cats_by='Factor', cats_order=None, col_wrap=3,
                              magnitude=False, rename=True, exclude_vars=None,
                              show=True, save=True, figsize=[26, 10], sharey=True,
                              legend_cols=None, legend_loc='lower center', legend_mode='expand',
                              fig_legend=True, ylim=None, p_annotate=True, **kwargs):
        """Plot box plots of correlation or regression coefficients."""
        if type(exclude_vars) is str:
            exclude_vars = [exclude_vars]  # ensure iterable
        if type(reverse_code) is str:
            reverse_code = [reverse_code]  # ensure iterable
        from functions_documentation import string_list
        keys = self.__getattribute__(attribute).keys()
        if (p_annotate) and ('_p' not in attribute):
            attribute = attribute + '_p'
        if dv_groups_keys is None:
            dv_groups_keys = keys
        if dv_groups is None:
            dv_groups = [self.__getattribute__(attribute)[k] for k in dv_groups_keys]
        dvs_grped = dict(zip(dv_groups_keys, dv_groups))
        models, sts = [self, comparison_model], '*' if self.alpha >= 0.05 else '**' if self.alpha >= 0.01 else '***'
        stars = sts.count('*')  # number of stars to be significant
        mat = [pd.concat([m.__getattribute__(attribute)[k] for k in keys], axis=1) for m in models]
        if fac_groups is not None:  # group factors together (if desired)
            mat = [pd.concat([m.loc[:, f, :] for f in fac_groups], keys=fac_groups_keys) for m in mat]
        if p_annotate:
            mat = [m.astype(str).applymap(lambda x: np.nan if x.count('*') < stars else x.strip('*')) for m in mat]
        if exclude_vars is not None:  # drop certain variables (if desired)
            mat = [m.drop(wh([e in m.columns for e in exclude_vars], 1, exclude_vars), axis=1) for m in mat]
        mat = [m.astype(float).abs() if magnitude else m.astype(float) for m in mat]  # abs() and/or float
        if reverse_code is not None:  # reverse magnitudes of certain variables (if desired)
            for m in range(len(mat)):
                for v in reverse_code:
                    if v in mat[m].columns:
                        mat[m].loc[:, v] = abs(mat[m].loc[:, v])
        if (rename) and ('rename' in self.dictionaries.keys()):  # use internal rename (if not provided)
            rename = self.dictionaries['rename']
        if rename is not None:
            mat = [m.rename(rename, axis=0).rename(rename, axis=1) for m in mat]
            dvs_grped = dict(zip(dvs_grped.keys(),
                                 [[rename[x] if x in rename.keys() else x for x in dvs_grped[k]] for k in dvs_grped]))
        matrices = pd.concat(mat, keys=model_keys)
        ix_names = [comparison_type] + [['Cluster'], []][fac_groups is None] + ['Model', 'Factor'] + ['Outcome']
        matrix = pd.DataFrame(pd.concat([matrices[dvs_grped[k]] for k in dvs_grped],
                                        axis=1, keys=dvs_grped.keys()).stack(1).stack(),
                              columns=[att_lab]).rename_axis(ix_names + ['Correlates']).reset_index()
        fig = sb.catplot(data=matrix, x=cats_by, y=att_lab, col='Correlates', col_wrap=col_wrap,
                         sharey=True, hue=ix_names[0], kind='boxen', order=cats_order, **kwargs)
        fig.set_axis_labels('', att_lab + ['', ' (Magnitude)'][magnitude is True])
        fig.set_titles('{col_name}')
        fig.despine(left=True)
        # min_corr, max_corr = min([m.min().min() for m in mat]), max([m.max().max() for m in mat])
        # fig, axes = plt.subplots(*square_grid(len(dv_groups_keys)), figsize=figsize, sharey=sharey)
        # for i, k in enumerate(dvs_grped):
        #     # df, df_comp = mat[0][dvs_grped[k]], mat[1][dvs_grped[k]] # results for models
        #     # if cats_by == ix_names[1]: # if grouping by rows...
        #     #     df.columns = ['%s\n%s'%(model_keys[0], c) for c in df.columns]
        #     #     df_comp.columns = ['%s\n%s'%(model_keys[1], c) for c in df_comp.columns]
        #     # else: # if grouping by columns (outcomes)
        #     #     df.columns = ['%s\n%s'%(model_keys[0], c) for c in df.columns]
        #     #     df_comp.columns = ['%s\n%s'%(model_keys[1], c) for c in df_comp.columns]
        #     cat_mat = pd.DataFrame(matrices[dvs_grped[k]].stack().rename_axis(ix_names), columns=[att_lab])
        #     if cats_order is not None:
        #         cat_mat = cat_mat.reset_index().set_index(cats_by).loc[cats_order].reset_index().set_index(ix_names)
        #     sb.boxenplot(data=cat_mat.reset_index(), x=cats_by, y=att_lab, hue=ix_names[0],
        #                  ax=axes.ravel()[i], **kwargs)
        #     axes.ravel()[i].set_ylabel(att_lab + ['', ' (Magnitude)'][magnitude])
        #     if ylim is not None: axes.ravel()[i].set_ylim(ylim)
        #     axes.ravel()[i].set_title(k)
        #     if fig_legend:
        #         axes.ravel()[i].legend_.remove()
        #     else:
        #         if legend_cols is None: legend_cols = len(pd.unique(cat_mat.reset_index()[ix_names[0]]))
        #         axes.ravel()[i].legend(ncol=legend_cols, loc=legend_loc, frameon=False)
        # for i, ax in enumerate(axes.ravel()):
        #     if i >= len(dvs_grped):
        #         ax.tick_params(which='minor', bottom=False, left=False)
        #         ax.grid(which='minor', color='w', linestyle='-', linewidth=3)
        #         for edge, spine in ax.spines.items(): # turn off spines & edges
        #             spine.set_visible(False)
        # if len(self.model_names) is None: legend_cols = len(self.model_names)
        # if fig_legend: # figure-level legend (if desired)
        #     fig.legend(*[q for q in axes.ravel()[i].get_legend_handles_labels()], loc=legend_loc,
        #                mode=legend_mode, ncol=legend_cols, frameon=False) # legend (horizontal)
        if show:
            plt.get_current_fig_manager().window.showMaximized()  # maximize plot window
            if fig_legend is False:
                fig.tight_layout()
            plt.show()
        else:
            plt.close('all')
        if save:
            caption = 'Boxplots of %s Model-Factors\' %s' % (self.analysis, att_lab)
            if self.files['file_suffix'] != 0:
                caption = caption + ' (' + self.files['file_suffix'] + ')'  # (suffix)
            if 'figures' not in self.files.keys():
                self.files.update({'figures': {}})
            if attribute not in self.captions.keys():
                self.captions.update({attribute: {}})
            if 'Plots' not in os.listdir():
                os.system('mkdir Plots')  # make Plots directory if absent
            # suf = '_' + self.files['file_suffix'] if self.files['file_suffix'] != 0 else ''
            file = 'Plots/Boxplots_%s_%s_%s_%s.png' % (attribute, cats_by, *[re.sub(' ', '_', i) for i in model_keys])
            self.files['figures'].update({'Boxplots_%s_%s' % (attribute, '%s_%s' % tuple(model_keys)): file})
            if exclude_vars is not None:
                caption = caption + ' (Excluding ' + string_list(exclude_vars) + ')'
            caption = re.sub('[)] [(]', '; ', caption)
            self.captions.update({'Boxplots_Comparison': caption})
            fig.savefig(file)
        return fig

    # Write Tables
    def export_tables_excel(self, attributes=None, rn_dict={}, digits=3,
                            file_stem=None, sig=None, captions=None, style='Academic'):
        """Export tables to Excel."""
        if 'dictionaries' not in dir(self):
            self.dictionaries = {'rename': rn_dict}
        if 'rename' not in list(self.dictionaries.keys()):
            self.dictionaries.update({'rename': rn_dict})
        rn_dict = self.dictionaries['rename']
        if attributes is None:
            attributes = wh(['results_' in x for x in dir(self)], True, dir(self))
        # flt_form = str('%%.%df'%digits)
        fsa = []
        if file_stem is None:
            file_stem = re.sub('_0', '', '%s_%s' % (self.analysis, self.files['file_suffix']))
        for a in attributes:  # for all results attributes
            if a in dir(self):
                tables = self.__getattribute__(a)
                if ((type(tables) == list) or (type(tables) == dict)) and (len(tables) > 1):  # if multiple tables...
                    file, write_mode = '%s_%s.xlsx' % (file_stem, a), 'w'
                    shs = list(tables.keys()) if type(tables) == dict else np.array(range(len(tables)))  # sheet names
                else:  # otherwise, put in "other" spreadsheet & name sheet after attribute
                    file = 'results_other_%s.xlsx' % file_stem
                    tables, shs = dict({a: tables}), [a]  # make dictionary out of table & sheet name = attribute name = key
                    # write_mode = 'a' if file in os.listdir('Results') else 'w' # append if already exists
                    write_mode = 'w'  # append if already exists
                print('Attempting to write %s to %s, sheets %s' % (a, 'Results/' + file, str(shs)))
                writer = pd.ExcelWriter('Results/' + file, mode=write_mode)
                for t in shs:
                    if pd.DataFrame(tables[t]).empty:
                        warnings.warn('\n%s\n\nCould not print table %s for %s\n\n%s' % ('*' * 80, t, a, '*' * 80))
                        continue
                    tb = tables[t].rename(rn_dict, axis=1).rename(rn_dict)  # "prettier" names
                    tb = tb.style.applymap(lambda x: color_tab(x, sig=sig))  # green = +, red = -, black = N.S. (with sig)
                    if captions is not None:
                        tb = tb.set_caption(self.captions[a][t])
                    tb = tb.set_table_styles(style)
                    tb.to_excel(writer, sheet_name=str(t), engine='openpyxl')
                writer.save()
                fsa = fsa + [(file, shs, a)]  # (files, sheets, attributes)
            else:
                warnings.warn(f'\n\n***** Attribute {a} not found!')
        if 'tables' not in list(self.files.keys()):
            self.files.update({'tables': {}})  # + tables key to files if absent
        self.files['tables'].update({'excel': dict(zip([f[2] for f in fsa], [f[0] for f in fsa]))})
        self.files['tables'].update({'excel_sheets': dict(zip([f[2] for f in fsa], [f[1] for f in fsa]))})

    def check_data(self, name='', precision=3):
        """Check data match (attribute, file, MPlus)."""
        file_mplus, file_data = self.files['data_mplus_file'], self.files['data_file']
        df = pd.read_csv(file_mplus, header=None, names=self.data_mplus.columns).set_index(self.mplus_index)
        df = df[pd.Series(self.DV_subsets_dictionary).explode()].replace('*', np.nan)  # subset columns; * -> NaN
        att = self.data.set_index(self.data_info['index_col'])[df.columns].astype(float)  # model data
        df = df.assign(CASEID=[att.index.values[int(i)] for i in df.index.values]).set_index('CASEID')
        df = df.rename({'CASEID': self.data_info['index_col']}, axis=1)
        file = pd.read_csv(file_data, index_col=self.data_info['index_col'])[df.columns].loc[att.index.values]
        [print(df.astype(float).round(precision).compare(x.round(precision))) for x in [att, file, df]]

    def check_tables_r(self, res_dir='Results', digits=3):
        """Compare to R results."""
        rn = self.dictionaries['rename']  # rename dictionary
        attributes = list(self.files['tables']['excel'].keys())  # attribute names
        comparison = dict(zip(attributes, [np.nan] * len(attributes)))
        difference = dict(zip(attributes, [np.nan] * len(attributes)))
        contrast = dict(zip(attributes, [np.nan] * len(attributes)))
        for a in attributes:  # iterate attributes
            file, shs = [self.files['tables'][k][a] for k in ['excel', 'excel_sheets']]  # file & sheet name
            tp = self.__getattribute__(a)
            if not isinstance(tp, dict):
                try:
                    tp = dict(zip(shs, [i.rename(rn, axis=1).rename(rn) for i in tp]))
                except Exception as err:
                    print(err)
                    warnings.warn('Comparison for attribute %s failed.' % a)
                    continue
            ix = np.arange(len(tp[shs[0]].index.names))
            if file not in os.listdir(res_dir):
                warnings.warn('%s not in %s directory' % (file, res_dir))
                continue
            tx = dict(zip(shs, [pd.read_excel('%s/%s' % (res_dir, file), s, index_col=ix) for s in shs]))
            if file in os.listdir('%s/R/' % res_dir):
                try:
                    f_r = '%s/R/%s' % (res_dir, file)  # R version of results
                    tr = dict(zip(shs, [pd.read_excel(f_r, s) for s in shs]))
                    tr = dict(zip(shs, [tr[s].set_index(list(tr[s].columns[:len(ix)])) for s in shs]))
                    tr = dict(zip(shs, [tr[s].rename_axis(tp[s].index.names) for s in shs]))
                    self.__setattr__(a + '_R', tr)  # set R version as attribute with <a>_R suffix
                except Exception as e:
                    print(e, '\n\n%s\n\nCould not extract R version of %s.\n\n%s' % ('*' * 80, a, '*' * 80))
            try:
                comp_x, diff_x, contr_x = compare_data(tp, tx)
                if all([d.round(digits).max().max() == 0 for d in diff_x]):
                    h = '\n\n%s\n\n' % ('=' * 80)
                    print('%s%s & Excel output line up (%d digits of precision)%s' % (h, a, digits, h))
                else:
                    warnings.warn('%sMISMATCH: Python attributes & Excel %s output (%d digits)%s' % (h, digits, a, h))
            except Exception as err:
                print(err)
                warnings.warn('Failed to compare to Excel output for attribute %s' % a)
            try:
                comp_r, diff_r, contr_r = compare_data(tp, tr, object_names=['Python', 'R'])
                comparison[a], difference[a], contrast[a] = comp_r, diff_r, contr_r
            except Exception as err:
                warnings.warn('Failed to compare to R output for attribute %s' % a)
                comparison.pop(a)
                difference.pop(a)
                contrast.pop(a)
                print(err)
        self.check_results = {'comparison': comparison, 'difference': difference, 'contrast': contrast}

    def export_tables_word(self, attributes, supplement_tables=False, open_files=True, out_file=None,
                           write_results=True, sep_vars=['PTSD'], supplement_include=True, main_include=True,
                           addl_models=None, addl_attributes=None, addl_supplement_tables=None,
                           headers=['Tables and Figures', 'Supplementary Tables'], addl_model_heads=None,
                           table_of_contents=False, table_of_contents_supplement=True, orient='landscape'):
        """Write Tables to Documents."""
        # from functions_documentation import write_table
        # mfs = dict(zip(self.model_names, self.fac_series))
        # content_dict = self.model_content.apply(lambda x: dict(zip(mfs[x.name], self.fac_content[x.name])),
        #                                         result_type='reduce')
        if 'table_numbers' not in dir(self):
            self.table_numbers = {}
        if type(supplement_tables) is bool:
            supplement_tables = [supplement_tables] * len(attributes)
        if type(attributes) is str:
            attributes = [attributes]
        if addl_models is not None:
            if type(addl_attributes) is str:
                addl_attributes = [addl_attributes]
            if len(addl_attributes) == 1:
                addl_attributes = addl_attributes * len(addl_models)
            if type(addl_supplement_tables) is str:
                addl_supplement_tables = [addl_supplement_tables]
            if len(addl_supplement_tables) == 1:
                addl_supplement_tables = addl_supplement_tables * len(addl_models)
            for m in range(len(addl_models)):
                if 'table_numbers' not in dir(addl_models[m]):
                    addl_models[m].table_numbers = {}
                if type(addl_supplement_tables[m]) is bool:
                    addl_supplement_tables[m] = [addl_supplement_tables[m]] * len(addl_attributes[m])
        # count, count_supp = self.table_start, self.table_start_supplement  # table number count
        models = [self] if addl_models is None else [self] + addl_models  # additional model tables
        supp_tabs = [supplement_tables] + addl_supplement_tables  # whether attribute table(s) in supplement (by models)
        mod_attribs = [attributes] + addl_attributes
        tables = [[models[m].__getattribute__(a) for a in mod_attribs[m]] for m in range(len(models))]
        mod_num_ts = [[len(i) if type(i) is dict else 1 for i in m] for m in tables]  # of tables/attribute ~ model
        mod_att = [functools.reduce(lambda i, j: i + j,
                                    [[i[1]] * i[0] for i in zip(m[0], m[1])]) for m in zip(mod_num_ts, mod_attribs)]
        # Table of Contents
        ts = [[models[m].__getattribute__(a) for a in mod_attribs[m]] for m in range(len(models))]
        nr = range(len(models))
        titles = [[models[m].captions[a] for a in ([attributes] + addl_attributes)[m]] for m in range(len(models))]
        sm = [[[i[1]] * len(i[0]) if type(i[0]) is dict else [i[1]]
               for i in zip(titles[m], supp_tabs[m])] for m in nr]
        tab_lists = [[dict_part(i, 'items') if type(i) is dict else [i] for i in t] for t in ts]
        title_lists = [[dict_part(i, 'items') if type(i) is dict else [i] for i in t] for t in titles]
        # tab_cap = [[functools.reduce(lambda i, j: i + j, y) for y in x] for x in [tab_lists, title_lists]]
        mod_tabs = [functools.reduce(lambda i, j: i + j, t) for t in tab_lists]
        mod_titles = [functools.reduce(lambda i, j: i + j, t) for t in title_lists]
        supp_l = [[[t[0]] * len(t[1]) if type(t[1]) is dict else [t[0]] for t in zip(supp_tabs[m], ts[m])]
                  for m in nr]
        mod_supp = [functools.reduce(lambda i, j: i + j, t) for t in supp_l]
        tab_d = [dict(zip(m[0], m[1])) for m in zip(mod_titles, mod_tabs)]  # list of dictionaries of titles & tables
        supp_d = [dict(zip(mod_titles[m], pd.Series(sm[m]).explode())) for m in range(len(models))]  # supp ~ title
        m_t_main, m_t_supp = [[wh(list(pd.Series(m)), x, list(m.keys())) for m in supp_d] for x in [False, True]]
        all_nums = {}
        if main_include:  # initialize TOC(s)
            main_titles = list(pd.Series(m_t_main).explode().dropna())
            main_nums = dict(zip(main_titles, ['%d' % (self.table_start + n) for n in range(len(main_titles))]))
            all_nums = {**all_nums, **main_nums}
            caps = ['Table %s. %s' % (main_nums[t], t) for t in main_nums]
            if type(caps) is str:
                caps = [caps]  # ensure iterable
            if table_of_contents:  # main TOC
                head_p = [headers[1] in p.text for p in self.main['Tables'].paragraphs]
                if any(head_p):
                    toc_loc = wh(head_p, True)
                    if sum(head_p) > 1:
                        toc_loc = toc_loc[0]
                    toc = self.main['Tables'].paragraphs[toc_loc]
                else:
                    toc = self.main['Tables'].add_paragraph()
                    head_line = toc.add_run(headers[1] + '\n')
                    head_line.bold = True
                    head_line.font.name = 'Times'
                    head_line.font.size = 12
                if self.table_start > 1:
                    self.main['Tables'].add_page_break()
                for c in caps:
                    line = toc.add_run('\n' + c)
                    line.font.name = 'Times'
                    line.font.size = 12
                if main_include:
                    self.main['Tables'].add_page_break()
        if supplement_include:
            supp_titles = list(pd.Series(m_t_supp).explode().dropna())
            st_r = range(len(supp_titles))
            s_nums = dict(zip(supp_titles, ['S%d' % (self.table_start_supplement + n) for n in st_r]))
            all_nums = {**all_nums, **s_nums}
            caps_supp = ['Table %s. %s' % (s_nums[t], t) for t in s_nums]
            if self.table_start_supplement > 1:
                self.supplement['Tables'].add_page_break()
            if table_of_contents_supplement:  # main TOC
                head_p_supp = [headers[1] in p.text for p in self.supplement['Tables'].paragraphs]
                if any(head_p_supp):
                    toc_supp_loc = wh(head_p_supp, True)
                    if sum(head_p_supp) > 1:
                        toc_supp_loc = toc_supp_loc[0]
                    toc_supp = self.supplement['Tables'].paragraphs[toc_supp_loc]
                else:
                    toc_supp = self.supplement['Tables'].add_paragraph()
                    head_line = toc_supp.add_run(headers[1] + '\n')
                    head_line.bold = True
                    head_line.font.name = 'Times'
                    head_line.font.size = 12
                if type(caps_supp) is str:
                    caps_supp = [caps_supp]  # ensure iterable
                for c in caps_supp:
                    line = toc_supp.add_run('\n' + c)
                    line.font.name = 'Times'
                    line.font.size = 12
                if supplement_include:
                    self.supplement['Tables'].add_page_break()
        # Store Table Numbers
        num_dict = [dict(zip(m, [all_nums[i] for i in m])) for m in mod_titles]
        table_numbers = []
        for m in range(len(models)):
            n, a, t = num_dict[m], mod_att[m], titles[m]
            d = [(a[wh(t, i)], i if type(i) is str else dict(zip(list(i.keys()), [n[i[k]] for k in i]))) for i in t]
            table_numbers = table_numbers + [dict(d)]
        self.table_numbers, self.table_numbers_addl = table_numbers[0], table_numbers[1:]
        # Write Tables
        kwargs_tab = {'note_dictionary': self.dictionaries['abbreviations'],
                      'rename_dictionary': self.dictionaries['rename']}
        rn = self.dictionaries['rename']
        # att_intersect = functools.reduce(lambda i, j: [v for v in i if v in j], [attributes] + addl_attributes)
        for m in range(len(models)):
            # tn = table_numbers[m] # table number dictionary
            for t in range(len(mod_titles[m])):
                ct = num_dict[m][mod_titles[m][t]]
                # tnk = list(tn.keys())
                tab, supp = tab_d[m][mod_titles[m][t]], mod_supp[m][t]
                # a = wh([ct in dict_part(tn[k], 'items') if type(tn[k]) is dict else k for k in tn], 1, tnk)
                a = wh([mod_titles[m][t] in q for q in title_lists[m]], 1, mod_attribs[m])
                n_e = '\n* p < 0.05. ** p < 0.01. *** p < 0.001.' if '_p' in mod_att[m][t] else None
                ort = orient if type(orient) is str else orient[m][a]
                if all([p.text == '' for p in [self.main['Tables'], self.supplement['Tables']][supp].paragraphs]):
                    sec = [self.main['Tables'], self.supplement['Tables']][supp].sections[-1]
                else:
                    sec = [self.main['Tables'], self.supplement['Tables']][supp].add_section(WD_SECTION.NEW_PAGE)
                sec.orientation = wh(['portrait', 'landscape'], ort)  # change to desired orientation
                sec.page_width = Inches([8.5, 11][wh(['portrait', 'landscape'], ort)])
                sec.page_height = Inches([11, 8.5][wh(['portrait', 'landscape'], ort)])
                file = self.documents[['main', 'supplement'][supp]]['Tables']
                print('Writing %s (%s) tables to %s.' % (a, mod_titles[m][t], file))
                tab.columns = [str(i) for i in tab.columns]  # make sure columns == strings
                tab.columns = [cap(rn[c]) if c in rn.keys() else cap(c) for c in tab.columns]  # re-name
                kws = {**kwargs_tab, 'blank_mid_index_border': True} if 'descriptives_outcomes' in a else kwargs_tab
                write_table(tab, self.documents[['main', 'supplement'][supp]]['Tables'],
                            note_extras=n_e, title=mod_titles[m][t], table_count=ct,
                            document=[self.main['Tables'], self.supplement['Tables']][supp], **kws)
                [self.main['Tables'], self.supplement['Tables']][supp].save(file)

    def export_figures_word(self, attributes=[[None], [None]], parts=['main', 'supplement'], addl_models=None,
                            wide=[Inches(7), Inches(4.74)], tall=[Inches(4.74), Inches(7)], note=None,
                            reinitialize=True, exclude_DV_subset=None):
        """Add Figures to Document."""
        models = [self] if addl_models is None else [self] + addl_models
        if type(addl_models) is not list:
            addl_models = [addl_models]
        for w in range(len(parts)):
            if attributes[w] is None:
                attributes[w] = list(self.files['figures'][parts[w]].keys())
        if type(attributes) is not list:
            attributes = [attributes]
        if type(attributes[0]) is not list:
            attributes = [[a] for a in attributes]
        # kws = dict(zip(['landscape', 'portrait'], [dict(zip(['width', 'height'], x)) for x in [wide, tall]]))
        kws = {'landscape': {'width': 8}, 'portrait': {'height': 8}}
        fig_num = 0
        for mod in models:
            print('Figures for %s %s' % (self.analysis, self.caption_suffix))
            for w in parts:
                # fig_sec = self.documents[w]['Figures'] if reinitialize else None
                doc = docx.Document()
                heading_par = doc.add_paragraph()
                heading = heading_par.add_run('Figures')
                heading.bold = True
                heading.font.name = 'Times New Roman'
                heading.font.size = 12
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                heading.line_spacing_rule = WD_LINE_SPACING.SINGLE
                atts, figs = attributes[wh(parts, w)], mod.files['figures']
                files = [f if type(f) is list else [(f)] for f in [figs[w][a] for a in atts]]  # iterable
                names, orients = [dict(zip(atts, [[i[x] for i in f] for f in files])) for x in [0, 1]]
                for a in atts:
                    for i in zip(names[a], orients[a]):
                        if exclude_DV_subset is not None:
                            if i[0] in exclude_DV_subset:
                                continue
                        fig_num += 1
                        if all([p.text == '' for p in doc.paragraphs]):
                            sec = doc.sections[-1]  # start at top/last if don't need to add page break
                        else:
                            sec = doc.add_section(WD_SECTION.NEW_PAGE)  # page break if not blank
                        sec.orientation = wh(['portrait', 'landscape'], i[1])  # change to desired orientation
                        sec.page_width = Inches([8.5, 11][wh(['portrait', 'landscape'], i[1])])
                        sec.page_height = Inches([11, 8.5][wh(['portrait', 'landscape'], i[1])])
                        head = doc.add_paragraph()
                        head.paragraph_format.keep_together = True
                        # head.paragraph_format.keep_with_next = True
                        head.paragraph_format.line_spacing = 0
                        fig_title = mod.captions[a]
                        if type(mod.captions[a]) is dict:
                            fig_title = fig_title[list(mod.captions[a].keys())[wh(names[a], i[0])]]
                        pic_fig = head.add_run('Figure %s\n' % str(fig_num))
                        pic_fig.bold = True
                        pic_fig.font.name = 'Times New Roman'
                        pic_fig.font.size = 12
                        pic_hd = head.add_run(fig_title)
                        pic_hd.italic = True
                        pic_hd.font.name = 'Times New Roman'
                        pic_hd.font.size = 12
                        doc.add_picture(i[0], **kws[i[1]])  # add figure
                        # pic_hd.add_picture(i[0], **kws[i[1]]) # add figure
                        # doc.add_picture(i[0], **kws[i[1]]) # add figure
                        if note is not None:
                            para = doc.add_paragraph()
                            nhd = para.add_run('\nNote. ')
                            nhd.italic = True
                            nhd.font.name = 'Times New Roman'
                            nhd.font.size = 10
                            if mod.dictionaries['abbreviations'] is not None:  # abbreviations for table note?
                                text = fig_title.split(' ')
                                tab = self.__getattribute__(a)  # table
                                if type(tab) is dict:
                                    tab = tab[list(mod.captions[a].keys())[wh(names[a], i[0])]]
                                text = text + list(pd.Series(tab.index.values).explode())  # add rows to table text list
                                text = text + list(pd.Series(tab.columns).explode())  # add columns to table text list
                                abbs = wh([k in text for k in mod.dictionaries['abbreviations']], 1,
                                          list(mod.dictionaries['abbreviations'].keys()))
                                abb_notes = ['%s = %s' % (a, mod.dictionaries['abbreviations'][a]) for a in abbs]
                                n = '%s %s.' % (note, functools.reduce(lambda x, y: f'{x}, {y}', abb_notes))
                                note_txt = para.add_run(n.strip())
                                note_txt.italic = False
                                note_txt.font.name = 'Times New Roman'
                                note_txt.font.size = 10
                            else:
                                note_txt = para.add_run(note.strip())
                                note_txt.font.name = 'Times New Roman'
                                note_txt.italic = False
                                note_txt.font.size = 10
                if w == 'supplement':
                    self.supplement['Figures'] = doc
                    self.supplement['Figures'].save(self.documents[w]['Figures'])
                else:
                    self.main['Figures'] = doc
                    self.main['Figures'].save(self.documents[w]['Figures'])

    def write_results(self, sep_vars=['PTSD'], addl_models=None, types=['fit', 'correlations', 'regressions'],
                      types_supplement=[False, False, True], models_supplement=False,
                      model_heads=None, sig_stars=None, star_p=[0.05, 0.01, 0.001], abbreviations=None,
                      exclude_DV_subset=None):
        """Write Results to Document."""
        headers = ['Structural Analyses', 'Correlations', 'Regressions']
        type_headers = dict(zip(['fit', 'correlations', 'regressions'], headers))
        models = [self] if addl_models is None else [self] + addl_models
        if type(models_supplement is bool):
            models_supplement = [models_supplement] * len(models)
        if model_heads is None:
            model_heads = ['Model %d' % m for m in range(len(models))]
        if sig_stars is None:
            sig_stars = ['*', '**', '***'][wh(star_p, self.alpha)]  # significance threshold
        table_nums = [self.table_numbers] if addl_models is None else [self.table_numbers] + self.table_numbers_addl
        if type(sep_vars) is str:
            sep_vars = [sep_vars]  # ensure list
        res = dict()
        for m in range(len(models)):
            abbs = models[m].dictionaries['abbreviations'] if abbreviations is None else abbreviations
            all_DVs = pd.Series(dict_part(models[m].DV_subsets_dictionary, 'items')).explode()
            seps = wh([s in all_DVs for s in sep_vars], 1, sep_vars)
            atts = list(table_nums[m].keys())
            att_cor = wh([False if type(i) in [list, tuple] else 'results_correlations' in i.lower() for i in atts],
                         1, atts)  # name of correlation attribute with a table number from export_tables_word
            if type(att_cor) is list:
                att_cor = att_cor[0]
            kwds = {'separate_variables': seps if len(seps) > 0 else None, 'types': types, 'abbs': abbs,
                    'analysis': models[m].analysis, 'hypos': models[m].dictionaries['hypotheses'], 'headers': headers}
            if 'fit' in types:
                if models[m].mplus_arguments['ESTIMATOR'].lower() == 'mlr':
                    fit_ix = 'BIC'
                elif models[m].mplus_arguments['ESTIMATOR'].lower() == 'wlsmv':
                    fit_ix = 'Chi-Squared'
                elif models[m].mplus_arguments['ESTIMATOR'].lower() == 'bayes':
                    fit_ix = 'Posterior-Predictive P-Value'
                else:
                    fit_ix = 'Fit'
                kwds.update({'results_fit': models[m].results_fit, 'fit_index': fit_ix})
            if 'correlations' in types:
                kwds.update({'results_cor': models[m].__getattribute__(att_cor), 'attribute_cor': att_cor})
            if 'regressions' in types:
                kwds.update({'results_reg': models[m].__getattribute__('results_regressions_p'),
                             'sig_stars': sig_stars, 'star_p': star_p})
            out = results_section_documents(models[m].model_content, table_nums[m], **kwds)
            res.update({model_heads[m]: {}})
            for o in out:
                res[model_heads[m]].update({o: out[o]})
        for supp in [False, True]:  # iterate through main and/or supplement
            doc = self.supplement['Results'] if supp else self.main['Results']
            main_types, supp_types = [wh(types_supplement, x, types) for x in [0, 1]]
            if type(main_types) is str:
                main_types = [main_types]
            if type(supp_types) is str:
                supp_types = [supp_types]
            for t in [main_types, supp_types][supp]:
                secs = [res[m][type_headers[t]] for m in model_heads]  # result type t sections for models (if in same part)
                header = doc.add_paragraph()
                # header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # header.line_spacing_rule = WD_LINE_SPACING.SINGLE
                head = header.add_run(cap(type_headers[t]))
                head.bold = True
                head.font.name = 'Times New Roman'
                head.font.size = 12
                for a in secs[0].keys():
                    txt = dict(zip(model_heads, [s[a] for s in secs]))
                    para = doc.add_paragraph()
                    para.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    graph = para.add_run(a)
                    graph.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    graph.bold = True
                    graph.italic = True
                    graph.font.size = 12
                    graph.font.name = 'Times New Roman'
                    for k in txt:
                        if exclude_DV_subset is not None:
                            if k in exclude_DV_subset:
                                continue
                        if functools.reduce(lambda a, b: a + ' ' + b, txt[k]).strip() == '':
                            continue  # continue if blank section
                        small_header = para.add_run('\n\t' + k + '. ')
                        small_header.bold = True
                        small_header.font.size = 12
                        small_header.font.name = 'Times New Roman'
                        small_header.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        pg = para.add_run(re.sub('_', ' ', functools.reduce(lambda a, b: a + ' ' + b, txt[k])))
                        pg.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        pg.font.size = 12
                        pg.font.name = 'Times New Roman'
            if supp:
                self.supplement['Results'] = doc
                doc.save(self.documents['supplement']['Results'])
            else:
                self.main['Results'] = doc
                doc.save(self.documents['main']['Results'])
        return res

    def compose_manuscript(self, out_files, margins=[1, 1, 1, 1], header_title=None,
                           font_name='Times New Roman', font_size=12):
        """
        Combine sections into one manuscript and/or supplement.

        Parameters
        ----------
        out_files : list
            List of two strings (or None, if one or the other is not to be included)
            indicating the desired output file for the main and/or supplemental manuscript.
        margins : list
            [left, right, upper, lower] in inches

        Returns
        -------
        None.

        """
        self.manuscript = dict()
        for f, part in enumerate(['main', 'supplement']):
            if out_files[f] is not None:
                heads, files = [[(k, self.documents[part][k])[i] for k in self.documents[part]] for i in [0, 1]]
                self.manuscript.update({part: combine_word_documents(files)})
                for s in range(len(self.manuscript[part].sections)):  # iterate sections
                    self.manuscript[part].sections[s].left_margin = Inches(margins[0])
                    self.manuscript[part].sections[s].right_margin = Inches(margins[1])
                    self.manuscript[part].sections[s].top_margin = Inches(margins[2])
                    self.manuscript[part].sections[s].bottom_margin = Inches(margins[3])
                self.manuscript[part].save(out_files[f])
