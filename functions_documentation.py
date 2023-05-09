#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# pylint: disable=line-too-long
"""
Created on Wed Dec  9 11:07:50 2020.

@author: ena
"""

import functools
import numpy as np
import pandas as pd
import re
import warnings
import docx
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from functions_mplus import *
from functions_data_cleaning import wh, dict_part
from functions_visualization import cap


def orient(doc, orient):
    """Orient Document: Modified from j4n7 stackoverflow.com/questions/31893557/python-docx-sections-page-orientation#37681150."""
    if type(orient) is str:
        orient = wh(['portrait', 'landscape'], orient.lower())  # to integer (ENA)
    current_section = doc.sections[-1]
    curr_orient = current_section.orientation
    if type(curr_orient) is str:
        curr_orient = wh(['portrait', 'landscape'], curr_orient.lower())  # to integer (ENA)
    if orient != curr_orient:  # detect whether need to change page dimensions (ENA)
        new_width, new_height = current_section.page_height, current_section.page_width
    else:
        new_width, new_height = current_section.page_width, current_section.page_height
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = orient  # change to desired orientation (ENA)
    new_section.page_width = new_width
    new_section.page_height = new_height
    return new_section


def set_cell_border(cell, **kwargs):
    """
    From MadisonTrash at https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx.

    Set cell`s border
    Usage:
    set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                    bottom={"sz": 12, "color": "#00FF00", "val": "single"},
                    start={"sz": 24, "val": "dashed", "shadow": "true"}, end={"sz": 12, "val": "dashed"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def combine_word_documents(files, headers=None):
    """Combine Word documents: from maerteijn https://stackoverflow.com/questions/24872527/combine-word-document-using-python-docx."""
    merged_document = docx.Document()
    for index, file in enumerate(files):
        sub_doc = docx.Document(file)
        if len(sub_doc.paragraphs) == 0:
            continue
        # Don't add a page break if you've reached the last file.
        if index < len(files) - 1:
            # sub_doc.add_page_break()
            if headers is not None:
                head = sub_doc.paragraphs[0].insert_paragraph_before()
                hd = head.add_run(headers[index])
                hd.bold = True
                hd.font.name = 'Times New Roman'
                hd.font.size = 12
                hd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
    return merged_document


def ordinal(x, words=True):
    """Make Ordinals (e.g. second) & Superlatives (e.g. second best), modified from Gareth's code on codegolf."""
    out = "%d%s" % (x, "tsnrhtdd"[(x / 10 % 10 != 1) * (x % 10 < 4) * x % 10::4])
    out = re.sub('1st', 'first', re.sub('2nd', 'second', re.sub('3rd', 'third', re.sub('4th', 'fourth', out))))
    out = re.sub('5th', 'fifth', re.sub('6th', 'sixth', re.sub('7th', 'seventh', re.sub('8th', 'eighth', out))))
    out = re.sub('9th', 'ninth', re.sub('10th', 'tenth', out))
    return out


def superlative(x, top=3, bottom=7):
    """Make superlatives (e.g., best, second best, worst)."""
    out = ordinal(x)
    ords = [ordinal(n) for n in range(1, top + 1)]
    ords_r = [ordinal(n) for n in range(top + 1, bottom + 1)]
    ords_r.reverse()
    for q in ords:
        out = re.sub(q, q + ' best', out)
    for q in range(len(ords_r)):
        out = re.sub(ords_r[q], str(ordinal(q + 1)) + ' worst', out)
    out = re.sub('first ', '', re.sub('1st ', '', out))
    return out


def string_list(strings, abb_dict=None):
    """Text for Lists (e.g. just 1 = the input, 2 = input 1 and input 2, more = comma-separated list with &)."""
    if type(strings) is str:
        strings = [strings]
    if abb_dict is not None:
        strings = [abb_dict[s] if s in abb_dict.keys() else s for s in strings]
    if len(strings) >= 3:
        out = '%s, %s, and %s' % (strings[0], functools.reduce(lambda a, b: a + ', ' + b, strings[1:-1]), strings[-1])
        return re.sub('  ', ' ', out)
    elif len(strings) == 2:
        return re.sub('  ', ' ', strings[0] + ' and ' + strings[1])
    else:
        try:
            return strings[0]
        except Exception:
            return str(strings).strip('[').strip(']')


def results_section(tabulate, attribute, hypos, table_nums, abbs=None,
                    separate_variables=None, analysis='CFA', title='Correlations'):
    """Create Results Section Text."""
    tabulate = dict(zip(tabulate.keys(),
                        [tabulate[k].apply(lambda x: x.astype(str).str.strip('*').astype(float)) for k in tabulate]))
    DV_subsets_dictionary = dict([(k, list(tabulate[k].columns)) for k in tabulate])  # infer DV subsets dictionary
    tbs = []
    att = re.sub('results_', '', re.sub('_p', '', re.sub('_cis', '', re.sub('_tables', '', attribute))))
    if separate_variables is not None:
        for k in tabulate:
            t = tabulate[k].copy()
            for q in separate_variables:
                if q in tabulate[k].columns:
                    t = t.drop(q, axis=1)
        tbs = tbs + [t]
    else:
        tbs = [tabulate[k] for k in tabulate]
    # mini, maxi = min([abs(k).min().min() for k in tbs]), max([abs(k).max().max() for k in tbs])
    # sep = '' if separate_variables is None else '(other than with %s)' % string_list(separate_variables, abb_dict=abbs)
    # head_info = (cap(att), sep, '{:.3f}'.format(maxi - mini))
    # head = ['%s %s differed in magnitude across models, factors, and outcomes by %s or less.'%head_info]
    # sections_results = dict({cap(title) + '\n\t': head})
    sections_results = dict()
    counter = dict()
    for y in DV_subsets_dictionary:
        results = []
        table, tn = tabulate[y].rename_axis(['Model', 'Factor']), table_nums[attribute][y]
        varb = y if y.isupper() else y.lower()
        if separate_variables is not None:  # drop variables to be described separately
            for q in separate_variables:
                if q in table.columns:
                    table = table.drop(q, axis=1)
                    vab = '%s other than %s' % (varb, string_list(separate_variables))
                else:
                    vab = '%s' % varb
        else:
            table, vab = tabulate[y].rename_axis(['Model', 'Factor']), varb
        desc_abs = table.applymap(abs).describe()
        ran = 'in magnitude by %s or less' % '{:.3f}'.format(desc_abs.loc['max'].max() - desc_abs.loc['min'].min())
        miny, maxy = '{:.3f}'.format(table[table >= 0].min().min()), '{:.3f}'.format(table[table >= 0].max().max())
        out = 'by %s or less' % '{:.3f}'.format(table.apply(lambda q: abs(q).max() - abs(q).min(), axis=1).max())
        mod = 'by %s or less' % '{:.3f}'.format(table.apply(lambda q: abs(q).max() - abs(q).min(), axis=0).max())
        minyn, maxyn = '{:.3f}'.format(table[table < 0].min().min()), '{:.3f}'.format(table[table < 0].max().max())
        if (table[table < 0].dropna(how='all').empty) and table[table >= 0].dropna(how='all').empty:  # if not both + & -
            ranges = '%s to %s' % (miny, maxy) if all(table >= 0) else '%s to %s' % (minyn, maxyn)
        else:
            ranges = '%s to %s (positive values) and %s to %s (negative values)' % (miny, maxy, minyn, maxyn)
        di = (analysis, att, vab, ran, ranges)
        results = results + ['%s model %s with %s varied %s, ranging from %s.' % di]
        y_num = wh(list(DV_subsets_dictionary.keys()), y) % 3
        pref = [f'Magnitudes of {att} differed across outcomes {out}',
                f'{cap(att)} magnitudes varied across outcomes {out}',
                f'Magnitudes of {att} with distinct outcomes varied {out}'][y_num]
        suff = [f'within-outcome, across models\' factors {mod}',
                f'within-outcome across factors {mod}',
                f'by factors within-outcome {mod}'][y_num]
        results = results + ['%s and %s.' % (pref, suff)]
        if separate_variables is not None:
            if any([q in DV_subsets_dictionary[y] for q in separate_variables]):
                for q in separate_variables:  # iterate through variables reported separately
                    desc_sep = tabulate[y].rename_axis(['Model', 'Factor'])[[q]].describe()
                    info_sep = (att, q, '{:.3f}'.format(max(desc_sep.loc['max'] - desc_sep.loc['min'])),
                                '{:.3f}'.format(min(desc_sep.loc['min'])),
                                '{:.3f}'.format(max(desc_sep.loc['max'])))
                    results = results + ['the magnitude of %s with %s varied by %s or less (%s to %s).' % info_sep]
        unexpected = dict()
        for x in DV_subsets_dictionary[y]:
            tab = tabulate[y][[x]].rename_axis(['Model', 'Factor'])
            model_names = pd.unique([i[0] for i in tab.index.values])  # model names
            name = abbs[x] if (abbs is not None) and (x in abbs.keys()) else x.lower()
            unexpect = {}
            if x in hypos[y].keys():
                signs = hypos[y][x]
                for m in model_names:
                    ix = tab[x].loc[m].index.intersection(signs.keys())
                    unexpect.update(dict([(m, f) if tab[x].loc[m].loc[f] * signs[f] < 0 else (m, []) for f in ix]))
                for k in list(unexpect.keys()):
                    if len(unexpect[k]) == 0:
                        unexpect.pop(k)  # remove empty
            unexpected.update({x: unexpect})
            grped = tab.reset_index(0).groupby('Model')
            mins, maxs = [['{:.3f}'.format(q) for q in grped.apply(fx)[x]] for fx in [min, max]]
            ran = '{:.3f}'.format(max([float(q[0]) - float(q[1]) for q in zip(maxs, mins)]))
            x_num = wh(DV_subsets_dictionary[y], x) % 2
            verb = ['differed', 'varied'][wh(DV_subsets_dictionary[y], x) % 2]
            phr = ['across factors', 'across factors within-model', 'among model factors'][x_num]
            inform = (cap(att), name, verb, phr, ran, min(mins), max(mins))
            results = results + ['%s for %s %s %s by %s or less (%s to %s).' % inform]
        unexp_vars = wh([len(unexpected[w]) > 0 for w in unexpected], 1, list(unexpected.keys()))  # unexpect effect DVs
        if len(unexp_vars) == 0:
            exp_str = ['All %s were in expected directions (Table %s).' % (att, str(tn)),
                       'All factors associated in expected directions (Table %s).' % str(tn),
                       'Outcomes associated with all factors in expected directions (Table %s).' % (str(tn))]
            results = [exp_str[wh(list(DV_subsets_dictionary.keys()), y) % 3]] + results
        else:
            if type(unexp_vars) is str:
                unexp_vars = [unexp_vars]  # make sure iterable
            string, unexp_ys = [], [abbs[q] if (abbs is not None) and (q in abbs.keys()) else q for q in unexp_vars]
            for q in unexp_vars:
                unexp_facs, un_list = pd.unique([unexpected[q][t] for t in unexpected[q]]), list(unexpected[q])
                unexp_mods = [wh(dict_part(unexpected[q], 'items'), f, list(unexpected[q].keys())) for f in unexp_facs]
                mods = zip([wh(un_list, f, list(unexpected[q].keys())) for f in unexp_facs], unexp_mods)
                ms, mss = ' model', ' models'
                mods = [str(f[1]) + ms if type(f[0]) is str else str(string_list(f[1])) + mss for f in mods]
                facs = [abbs[f] if f in abbs.keys() else f for f in unexp_facs]
                string = string + [['the %s factor in the %s' % i for i in zip(facs, mods)]]
            string = string_list(['%s with %s' % w for w in zip(unexp_ys, [string_list(s) for s in string])])
            close = (att, string, str(tn))
            exp_phr = ['Some %s (%s) were in unexpected directions. (Table %s)' % close]
            results = exp_phr + results
        counter.update({y: unexpected})
        sections_results.update({cap(vab): results})
    sections = [q + ' ' + functools.reduce(lambda i, j: '%s %s' % (i, j), sections_results[q]) for q in sections_results]
    red = functools.reduce(lambda i, j: re.sub('variables outcomes', 'outcomes', i + ' ' + j), sections)
    results_text = re.sub('_', ' ', '\t' + red)
    return results_text, sections_results


def results_section_regression(results_regressions_p, hypos,
                               sig_stars=None, abbs=None, star_p=[0.05, 0.01, 0.001]):
    """Write Regression-Specific Results."""
    def mod_phrase(sig, num, y, mod_facs):
        ns = sig is False  # N.S.
        if len(mod_facs) > 0:
            phs = ['%s %s significantly associate%s with %s' % (cap(y), ['', 'did not'][ns], ['d', ''][ns], mod_facs),
                   '%s had %ssignificant associations with %s' % (mod_facs, ['', 'in'][ns], cap(y)),
                   '%s did %ssignificantly associate with %s' % (mod_facs, ['', 'not '][ns], cap(y)),
                   '%s had %ssignificant associations with %s' % (cap(y), ['', 'in'][ns], mod_facs),
                   '%s\'s associations with %s were %ssignificant' % (cap(y), mod_facs, ['', 'not '][ns])]
            out = phs[num % len(phs)]  # so will cycle through options (e.g. if num = 5, 1st option chosen)
        else:
            out = ''
        return out
    regs, hypos = results_regressions_p, hypos  # results & hypotheses dictionary
    text = {}
    for r in regs:
        sig = regs[r].applymap(lambda x: x.count('*') >= sig_stars.count('*'))  # significance
        # non_sig = dict(zip(sig.columns, [sig[x][sig[x] == False].index.values for x in sig.columns])) # N.S.
        paragraph = dict()
        for s in [True, False]:
            par = dict()
            for h in hypos[r]:  # iterate outcomes in hypotheses
                h0_f = list(hypos[r][h].keys())  # factors expected to have significant effects
                if h not in sig.columns:
                    continue
                # h0_f_ns = wh([k in [i[1] for i in non_sig[h]] for k in h0_f], 1, h0_f) # N.S. factors
                h0 = dict(zip(h0_f, [sig[h].loc[:, f][sig[h].loc[:, f] == s].index.values for f in h0_f]))  # N.S./S.
                facs = wh([len(h0[i]) > 0 for i in h0], 1, list(h0.keys()))  # keys for factors with models where N.S./S.
                if len(facs) == 0:
                    continue
                if type(facs) is str:
                    facs = [facs]  # make sure iterable
                h0 = dict(zip(facs, [h0[f] for f in facs]))  # just where N.S. or significant
                if s:
                    for f in h0:
                        q = regs[r][h].str.strip('*')
                        io = [(i, q.loc[i, f], star_p[regs[r].loc[i, f][h].count('*') - 1]) for i in h0[f]]
                        bp = ['%s (b = %s, p < %s)' % (i[0], i[1], '{:.3f}'.format(i[2])) for i in io]
                        h0.update({f: bp})
                strings = ['%s in the %s model%s' % (f, string_list(h0[f]), ['', 's'][len(h0[f]) > 1]) for f in h0]
                if '' in strings:
                    strings.remove('')
                phr = string_list(strings)
                to_replace = wh([k in h0_f + list(regs[r].keys()) for k in abbs.keys()], 1, list(abbs.keys()))
                to_replace = [to_replace] if type(to_replace) is str else to_replace
                for a in to_replace:  # re-name abbreviated factors
                    phr = re.sub(' %s ' % a, ' %s ' % abbs[a], ' ' + phr).strip(' ')
                phrase = mod_phrase(s, wh(list(hypos[r].keys()), h), h, phr)  # full phrase
                for a in abbs:
                    for ab in [' ' + a + ' ', '^' + a + ' ', ' ' + a + '$']:
                        if ab in phrase:
                            phrase = re.sub('  ', '', re.sub(ab, ' ' + abbs[a] + ' ', phrase)).strip()
                par.update({h: re.sub('_', ' ', cap(phrase).strip(' ') + '.')})  # update dictionary
            if len(par) == 0:
                continue
            else:
                paragraph.update({['NS', 'S'][s]: functools.reduce(lambda i, j: i + ' ' + j, pd.Series(par))})
        pag = []
        if 'S' in paragraph.keys():
            pag = pag + [paragraph['S']]
        if 'NS' in paragraph.keys():
            pag = pag + [paragraph['NS']]
        # text = text + [functools.reduce(lambda i, j: i + j, pd.Series(paragraph))]
        text.update({r: pag})
    # for a in list(pd.Series([list(regs[r].columns) for r in regs]).explode()): # replace more abbreviations
    #     if a in abbs.keys(): text = re.sub(' ' + a + ' ', ' ' + abbs[a] + ' ', re.sub('[.]', ' .', ' ' + text))
    # text = re.sub(' +', ' ', text).strip(' ')
    # print(text)
    return text


def results_section_fit(model_content, results_fit, fit_index='BIC', top=None):
    """Generate Fit Documentation."""
    if fit_index not in results_fit.columns:
        if fit_index in results_fit.index.values:
            results_fit = results_fit.T  # transpose if needed
        else:
            warnings.warn('%s not in table. Returning blank fit section.' % fit_index)
            return ''
    else:
        # Drop Models with Non-Comparable Indicator Sets from Fit
        mods = list(results_fit.index.values)  # models
        if top is None:
            top = model_content.shape[0]  # describe all models as best through mth best (instead of some as worst)
        nc = [model_content[[c]].replace('', np.nan).dropna().shape[0] < model_content.shape[0] for c in mods]
        if any(nc):  # if any non-comparable models, drop
            non_comp = [wh(nc, 1, mods)] if type(wh(nc, 1, mods)) is str else wh(nc, 1, mods)
            for x in non_comp:
                results_fit = results_fit.drop(x, axis=0)
        fit_cols_drop = wh(['Parameters' in c for c in results_fit.columns], 1, list(results_fit.columns))
        if (len(fit_cols_drop) == 0) or (type(fit_cols_drop) is str):
            fit_cols_drop = [fit_cols_drop]
        if 'N' in results_fit.columns:
            fit_cols_drop = fit_cols_drop + ['N']
        if len(fit_cols_drop) > 0:
            results_fit = results_fit.drop(fit_cols_drop, axis=1)
        results_fit = results_fit.dropna()
        ic_lists = [results_fit.sort_values(i) for i in results_fit.columns]  # sorted by values on fit indices
        for x in ['LL']:  # reverse for indices where lower is better
            if x in results_fit.columns:
                ix_ll = list(ic_lists[wh(results_fit.columns, x)].index.values)
                ix_ll.reverse()
                ic_lists[wh(results_fit.columns, x)] = ic_lists[wh(results_fit.columns, x)].loc[ix_ll]  # reverse
        sorted_mods = [list(i.index.values) for i in ic_lists]  # list of model names sorted on each fit index
        if any(pd.DataFrame(sorted_mods).apply(lambda x: len(pd.unique(x))) > 1):
            warnings.warn('Not all fit indices suggest the same ordering of models.')
        mod_sort = sorted_mods[wh(results_fit.columns, fit_index)]  # on specified fit index
        fits = results_fit.apply(
            lambda m: '(%s)' % string_list(['%s=%s' % (i[0], str(i[1])) for i in zip(results_fit.columns, m)]), axis=1)
        fits = fits.apply(lambda x: re.sub('and ', '', x))
        kw = {'top': top, 'bottom': len(mod_sort)}  # range for superlatives (e.g. third best-fitting)
        str_tups = [(mod_sort[m], superlative(m + 1, **kw)) for m in range(len(mod_sort))]
        first_third = dict(fits.loc[[i[0] for i in str_tups]][:kw['top']])
        first_third = [i + ' model ' + first_third[i] for i in first_third]
        fit_info = (first_third[0], string_list(first_third[1:]))
        text_fit = 'The %s was the best-fitting model, followed by the %s.' % fit_info  # best-fitting models
        # kms = list(dict(str_tups).keys())
        # rest = wh([i in [re.sub(' model.*', '', x) for x in first_third] for i in dict(str_tups)], 0, kms)  # the rest
        # cards = string_list([ordinal(wh(mod_sort, m) + 1, **kw) for m in rest])
        # text_fit = text_fit + ' The %s models were the %s-fitting models, respectively.'%(string_list(rest), cards)
        if any(nc):  # add sentence about non-comparable model(s) if needed
            info_nc = '%s model%s' % (string_list(non_comp), ['', 's'][sum(nc) > 1])
            nc_txt = 'The %s could not be compared to other models using fit indices,' % info_nc
            nc_txt = nc_txt + ' as %s different factor indicators.' % ['it includes', 'they include'][sum(nc) > 1]
            text_fit = re.sub(' +', ' ', text_fit + ' ' + nc_txt)
        return text_fit


def results_section_documents(model_content, table_numbers, hypos, abbs=None,
                              attribute_cor='results_correlations',
                              headers=['Structural Analyses', 'Correlations', 'Regressions'],
                              results_fit=None, results_cor=None, results_reg=None,
                              separate_variables=None, title='Results', analysis='CFA',
                              fit_index='BIC', types=['fit', 'correlations', 'regressions'],
                              sig_stars=None, star_p=[0.05, 0.01, 0.001]):
    """Write Results Document."""
    # doc.add_heading(title, 2)
    # types = ['fit', 'correlations', 'regressions']
    secs = dict()  # list for sub-sections
    # Fit
    if 'fit' in types:
        text_fit = results_section_fit(model_content, results_fit, fit_index=fit_index)
        secs.update({headers[0]: {'Fit': [text_fit]}})
    # Correlations
    if 'correlations' in types:
        res_kws = {'separate_variables': separate_variables, 'abbs': abbs, 'title': ''}
        text_cor = results_section(results_cor, attribute_cor, hypos,
                                   table_numbers, analysis=analysis, **res_kws)  # correlations
        secs.update({headers[1]: text_cor[1]})
    # Regressions
    if 'regressions' in types:
        text_reg = results_section_regression(results_reg, hypos, sig_stars=sig_stars, star_p=star_p, abbs=abbs)
        secs.update({headers[2]: text_reg})
    return secs


def create_table_note(table, abb_dict, key=None):
    """Table Notes."""
    labs = np.append(np.unique(pd.Series(table.index.to_list()).explode()), key)
    for c in table.columns:
        labs = np.append(labs, c)
    cell_labs = table.apply(lambda x: np.array(x) if type(x.iloc[0]) == str else None, result_type='reduce')
    for x in pd.unique(list(cell_labs.explode())):
        if x is not None:
            labs = np.append(labs, x)
    labs_dict = dict(zip(labs, [abb_dict.get(a) for a in labs]))  # abbreviations keys
    notes_list = ['' if labs_dict[k] is None else k + '=' + labs_dict[k] + '. ' for k in labs_dict]  # label = ''
    note = functools.reduce(lambda x, y: x + y, notes_list)  # put notes together
    return note


def write_table(matrix, file_name, title='', document=None, table_count=1, orientation=None,
                table_num_prefix='', note_dictionary=None, rename_dictionary={},
                note_extras=None, key=None, blank_mid_index_border=False, digits=3):
    """Write Tables to Word."""
    if document is None:
        document = docx.Document()  # start new empty document if none specified
    matrix = matrix.rename(rename_dictionary, axis=1).rename(rename_dictionary)  # rename
    matrix = matrix.replace(np.nan, '')
    # tables = document.sections[-1] # start adding tables at end of document
    table_section = document.sections[-1]
    if orientation is not None:
        if orientation.lower() == 'landscape':
            table_section.orientation = WD_ORIENT.LANDSCAPE
    else:
        table_section.orientation = WD_ORIENT.PORTRAIT
    label = document.add_paragraph()
    cap1 = label.add_run('Table %s%s.\n' % (table_num_prefix, re.sub('Table ', '', str(table_count))))
    cap1.bold = True
    cap1.font.name = 'Times New Roman'
    cap1.font.size = 12
    title_font = document.paragraphs[0].runs[0].font
    title_font.name = 'Times New Roman'
    title_font.size = 12
    # if title is not None:
    #     label.add_run('\n' + str(title) + '\n').italic = True
    if title is not None:
        label.add_run(str(title)).italic = True
    addl_rows = 0  # because doing title as text instead of through commented-out code for incorporating into table itself
    rows = range(np.shape(matrix)[0] + addl_rows)  # number of subsequent rows = # of data rows
    rows_init = 2 if type(matrix.columns) == pd.core.indexes.multi.MultiIndex else 1  # 1 or 2 headers
    max_span = matrix.shape[1] + len(matrix.index.names)  # maximum number of columns needed
    table = document.add_table(rows_init, max_span)  # start with 1 row X # columns + # indices
    table.style.font.name = 'Times New Roman'  # table font
    for r in range(len(table.rows)):  # iterate through headers
        headers = list(matrix.index.names) + list(matrix.columns)  # table headers (names of indices & columns)
        headers = [h[r] if type(h) == tuple else h for h in headers]  # get right level of headers (in case multi-header)
        headers = ['' if h is None else h for h in headers]  # replace None with blank in headers if needed
        cells = table.rows[r].cells  # initialize row 1 cells
        to_merge = [wh(headers, h) for h in pd.unique(headers)]  # list of arrays describing ranges of spanners
        # to_merge = wh([m.shape != () for m in to_merge], 1, to_merge) # only multi-cells in to_merge
        unique_header_1st_cells = [m[0] if type(m) is np.ndarray else m for m in to_merge]
        for c in range(len(headers)):  # iterate through headers
            if c >= len(matrix.index.values):  # borders if needed
                set_cell_border(cells[c], bottom={'sz': 10, 'val': 'single'}, top={'sz': 10, 'val': 'single'})
            if r <= rows_init:
                set_cell_border(cells[c], top={'sz': 10, 'val': 'single'})  # border for header row cells
            if (rows_init == 2) and (r == 0) and (c > len(matrix.index.names)) and (c not in unique_header_1st_cells):
                cells[c].text = ''  # blank if going to merge with previous spanner cells with shared heading
            else:
                cells[c].text = headers[c]  # initialize blank cells for title row columns
        if (rows_init == 2) and (r == 0):  # merge cells if multi-header
            for s in to_merge:
                try:
                    print(table.cell(0, s[0]).text)
                    print(table.cell(0, s[1]).text)
                    if len(s) > 1:
                        table.cell(0, s[0]).merge(table.cell(0, s[1]))
                except Exception:
                    pass
    unique_top_ixs = pd.unique([i[0] for i in matrix.index])  # unique indices
    unique_ix_occur1 = [np.where([i in j for j in matrix.index.values])[0][0] for i in unique_top_ixs]  # ix 1st occurs
    for r in rows:  # for subsequent rows (title, header, & actual results table rows)
        row_cells = table.add_row().cells
        for c in range(len(headers)):  # for each column
            if c >= len(matrix.index.names):  # if not in index columns
                value = matrix.iloc[r - addl_rows, c - len(matrix.index.names)]  # value (numerical or string, raw)
                if type(value) == str:
                    row_cells[c].text = value
                else:
                    row_cells[c].text = ('{:.%df}' % digits).format(value)  # format
                if blank_mid_index_border:  # in case there's a mid-table header...
                    if all([q == '' for q in matrix.index.values[r]]):  # detect if mid-table header
                        set_cell_border(row_cells[c], bottom={'sz': 10, 'val': 'single'},
                                        top={'sz': 10, 'val': 'single'})  # add border
            else:  # if index columns
                if (c == 0) and (len(matrix.index.names) > 1):  # if top level index column in multi-index
                    if r - addl_rows in unique_ix_occur1:  # if 1st row that a new top level index occurs, write to cell
                        row_cells[c].text = re.sub('_', ' ', matrix.index[r - addl_rows][0])
                else:  # otherwise, write successive indices in multi-index on every row
                    if len(matrix.index.names) > 1:  # if lower-level index in multi-index...
                        row_cells[c].text = matrix.index[r - addl_rows][c]  # write cth index in multi-index
                    else:  # if not multi-index...
                        row_cells[c].text = matrix.index[r - addl_rows]  # write cth index in multi-index
            if r == rows[-1]:
                set_cell_border(row_cells[c], bottom={'sz': 10, 'val': 'single'})  # bottom border
    if (note_dictionary is not None) or (note_extras is not None):  # add row for table note if needed
        table_note = '' if note_dictionary is None else create_table_note(matrix, note_dictionary, key=key)  # table note
        if note_extras is not None:
            table_note = table_note + ' ' + note_extras  # add any extra notes
        if table_note.strip('') != '':
            note = document.add_paragraph()
            pref = note.add_run('\nNote. ')
            pref.italic = True
            pref.font.name = 'Times'
            pref.font.size = 10
            note.runs[0].font.size = 10
            nt = note.add_run(table_note)
            nt.font.name = 'Times'
            nt.font.size = 10
            note.runs[1].font.size = 10
    document.add_page_break()
    document.save(file_name)
    return document
